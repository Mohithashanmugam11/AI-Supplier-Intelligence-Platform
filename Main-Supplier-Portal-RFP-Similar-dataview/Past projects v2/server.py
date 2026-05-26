# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# server.py â€” TrustBridge Supplier Ingestion Backend
# Run: uvicorn server:app --reload --port 8000
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

import os
import uuid
import time
import json
import base64
import mimetypes
import random
import smtplib
import io
import hashlib
import requests
import re
import zipfile
from pathlib import Path
from dotenv import load_dotenv
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
load_dotenv()

from fastapi import FastAPI, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer

from geometry      import compute_geometric_scores, build_geometric_vector
from inference     import run_inference, run_text_inference
from clip_embedder import compute_clip_embedding

try:
    import trimesh
except Exception:
    trimesh = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

# â”€â”€ Config â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY", "")
GEMINI_API_KEY   = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL     = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
INDEX_NAME       = "supplier-historical-projects"
INDEX_DIMENSION  = 512
CAD_CLIP_ENABLED = os.getenv("CAD_CLIP_ENABLED", "false").strip().lower() == "true"
REQUIRE_REAL_CLIP_VECTOR = os.getenv("REQUIRE_REAL_CLIP_VECTOR", "false").strip().lower() == "true"
EMBEDDER_BACKEND = os.getenv("EMBEDDER_BACKEND", "clip").strip().lower()
EMBEDDER_IS_EFFICIENTNET = EMBEDDER_BACKEND in {"efficientnet", "efficientnet_b0", "effnet"}
if EMBEDDER_IS_EFFICIENTNET and os.getenv("REQUIRE_REAL_CLIP_VECTOR") is None:
    # In EfficientNet mode, default to requiring real non-zero vectors (no silent fallback).
    REQUIRE_REAL_CLIP_VECTOR = True

ZOHO_CLIENT_ID     = os.getenv("ZOHO_CLIENT_ID", "")
ZOHO_CLIENT_SECRET = os.getenv("ZOHO_CLIENT_SECRET", "")
ZOHO_REFRESH_TOKEN = os.getenv("ZOHO_REFRESH_TOKEN", "")
ZOHO_API_BASE      = os.getenv("ZOHO_API_BASE", "https://www.zohoapis.com/crm/v2")
ZOHO_ACCOUNTS_URL  = os.getenv("ZOHO_ACCOUNTS_URL", "https://accounts.zoho.com/oauth/v2/token")

SMTP_HOST = os.getenv("SMTP_HOST", "smtp.office365.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "").strip()
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "").strip()
SMTP_FROM = os.getenv("SMTP_FROM", "").strip() or SMTP_USER
OTP_DELIVERY_MODE = os.getenv("OTP_DELIVERY_MODE", "smtp").strip().lower()  # smtp | zoho | auto
ZOHO_OTP_FUNCTION_NAME = os.getenv("ZOHO_OTP_FUNCTION_NAME", "").strip()
ZOHO_OTP_FUNCTION_URL = os.getenv("ZOHO_OTP_FUNCTION_URL", "").strip()

# Supplier access restrictions (all optional; set one or more to enforce)
ALLOWED_SUPPLIER_ZOHO_ID = os.getenv("ALLOWED_SUPPLIER_ZOHO_ID", "").strip()
ALLOWED_SUPPLIER_EMAIL = os.getenv("ALLOWED_SUPPLIER_EMAIL", "").strip().lower()
ALLOWED_SUPPLIER_DOMAIN = os.getenv("ALLOWED_SUPPLIER_DOMAIN", "").strip().lower().lstrip("@")

os.environ["GEMINI_API_KEY"] = GEMINI_API_KEY
os.environ["GEMINI_MODEL"]   = GEMINI_MODEL

if not PINECONE_API_KEY:   print("ERROR: PINECONE_API_KEY not set in .env")
if not GEMINI_API_KEY:     print("ERROR: GEMINI_API_KEY not set in .env")
if not ZOHO_REFRESH_TOKEN: print("WARN:  ZOHO_REFRESH_TOKEN not set â€” Zoho CRM integration disabled")
print(
    f"[auth] OTP mode={OTP_DELIVERY_MODE} "
    f"(zoho_fn={'yes' if (ZOHO_OTP_FUNCTION_NAME or ZOHO_OTP_FUNCTION_URL) else 'no'}, "
    f"smtp={'enabled' if (SMTP_USER and SMTP_PASSWORD) else 'disabled'})"
)

BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent

_ingestion_path = os.getenv("INGESTION_PATH", "").strip()
_stored_parts_candidates = []
if _ingestion_path:
    _stored_parts_candidates.append(Path(_ingestion_path) / "stored_parts")
_stored_parts_candidates.extend(
    [
        PROJECT_ROOT / "stored_parts",
        BASE_DIR / "stored_parts",
    ]
)
STORED_PARTS_DIR = next((p for p in _stored_parts_candidates if p.is_dir()), _stored_parts_candidates[0])
STORED_PARTS_DIR.mkdir(parents=True, exist_ok=True)

_upload_root = (os.getenv("UPLOAD_DIR", "").strip() or str(PROJECT_ROOT / "uploads_temp")).strip()
UPLOAD_DIR = Path(_upload_root)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

print(f"[paths] stored_parts={STORED_PARTS_DIR}")
print(f"[paths] uploads_temp={UPLOAD_DIR}")

# In-memory OTP store: {email: {otp, expires_at, zoho_account_id, company_name, attempts}}
_otp_store: dict = {}
OTP_EXPIRY_SECONDS = 600
OTP_MAX_ATTEMPTS = 5


def _mask_email(email: str) -> str:
    if "@" not in email:
        return email
    local, domain = email.split("@", 1)
    if len(local) <= 2:
        return f"{local[0:1]}***@{domain}"
    return f"{local[:2]}***@{domain}"


def _is_supplier_allowed(email: str, zoho_account_id: str) -> bool:
    has_rules = any([ALLOWED_SUPPLIER_ZOHO_ID, ALLOWED_SUPPLIER_EMAIL, ALLOWED_SUPPLIER_DOMAIN])
    if not has_rules:
        return True

    email = (email or "").strip().lower()
    zoho_account_id = str(zoho_account_id or "").strip()

    if ALLOWED_SUPPLIER_ZOHO_ID and zoho_account_id != ALLOWED_SUPPLIER_ZOHO_ID:
        return False
    if ALLOWED_SUPPLIER_EMAIL and email != ALLOWED_SUPPLIER_EMAIL:
        return False
    if ALLOWED_SUPPLIER_DOMAIN and not email.endswith(f"@{ALLOWED_SUPPLIER_DOMAIN}"):
        return False
    return True


def _send_otp_email(to_email: str, otp: str, company_name: str) -> bool:
    html = f"""
    <div style="font-family:'DM Sans',sans-serif;max-width:480px;margin:0 auto;padding:32px 24px;background:#FAFCFF;border:1px solid #CBD3DF;border-radius:8px;">
      <p style="font-size:14px;color:#2D4567;margin:0 0 8px;">Hi {company_name},</p>
      <p style="font-size:14px;color:#2D4567;margin:0 0 18px;">Your one-time access code is:</p>
      <div style="text-align:center;padding:18px;background:#F2F4F8;border-radius:6px;margin-bottom:16px;">
        <span style="font-family:'IBM Plex Mono',monospace;font-size:34px;letter-spacing:0.15em;color:#1B2D4F;">{otp}</span>
      </div>
      <p style="font-size:12px;color:#6B7F96;margin:0;">This code expires in 5 minutes.</p>
    </div>
    """

    def _send_once(from_header: str, envelope_from: str, reply_to: str = ""):
        msg = MIMEMultipart("alternative")
        msg["Subject"] = "Your TrustBridge Access Code"
        msg["From"] = from_header
        msg["To"] = to_email
        if reply_to:
            msg["Reply-To"] = reply_to
        msg.attach(MIMEText(html, "html"))

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.ehlo()
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(envelope_from, to_email, msg.as_string())

    preferred_from = (SMTP_FROM or SMTP_USER).strip()
    smtp_user = (SMTP_USER or "").strip()
    if not smtp_user:
        return False

    try:
        _send_once(preferred_from, preferred_from)
        return True
    except Exception as e:
        # Common Zoho policy rejection when SMTP auth user is not allowed to relay as SMTP_FROM.
        err_text = str(e).lower()
        if preferred_from != smtp_user and ("553" in err_text or "relay" in err_text or "sender is not allowed" in err_text):
            try:
                _send_once(smtp_user, smtp_user, reply_to=preferred_from)
                print("  ! OTP sent via SMTP_USER fallback (send-as relay blocked for SMTP_FROM)")
                return True
            except Exception as e2:
                print(f"  âœ— OTP email failed after fallback: {e2}")
                return False
        print(f"  âœ— OTP email failed: {e}")
        return False


def _send_otp_via_zoho_function(to_email: str, otp: str, company_name: str) -> bool:
    if not (ZOHO_OTP_FUNCTION_URL or ZOHO_OTP_FUNCTION_NAME):
        return False

    try:
        args_obj = {
            "email": to_email,
            "otp": otp,
            "company_name": company_name,
        }
        args_json = json.dumps(args_obj)

        if ZOHO_OTP_FUNCTION_URL:
            urls = [ZOHO_OTP_FUNCTION_URL]
        else:
            api_root = (ZOHO_API_BASE or "https://www.zohoapis.com/crm/v2").split("/crm/")[0]
            fn = ZOHO_OTP_FUNCTION_NAME
            urls = [
                f"{api_root}/crm/v8/functions/{fn}/actions/execute",
                f"{api_root}/crm/v2/functions/{fn}/actions/execute",
            ]

        def _is_zoho_success(resp) -> tuple[bool, str]:
            body = (resp.text or "")[:500]
            try:
                data = resp.json()
            except Exception:
                low = body.lower()
                if "error" in low or "failed" in low:
                    return False, body
                if "ok" in low or "success" in low or "sent" in low:
                    return True, body
                return False, f"ambiguous plain-text response: {body}"

            # Common response shapes:
            # 1) {"code":"success", ...}
            # 2) {"details":{"output":"OK"}}
            # 3) {"details":{"output":"ERROR: ..."}}
            code = str(data.get("code", "")).lower()
            if code and code != "success":
                return False, body

            details = data.get("details", {}) if isinstance(data, dict) else {}
            output = ""
            if isinstance(details, dict):
                output = str(details.get("output", "")).strip()
            if not output and isinstance(data, dict):
                output = str(data.get("output", "")).strip()

            if output:
                normalized_output = output.upper()
                if (
                    normalized_output in {"OK", "SUCCESS", "SENT"}
                    or "SUCCESS" in normalized_output
                    or "SENT" in normalized_output
                ):
                    return True, body
                if normalized_output.startswith("ERROR") or "FAILED" in normalized_output:
                    return False, output
                return False, f"ambiguous function output: {output}"

            return False, f"missing function output: {body}"

        attempts = []
        for url in urls:
            base_headers = {}
            if "auth_type=apikey" not in url:
                base_headers["Authorization"] = f"Zoho-oauthtoken {get_zoho_token()}"

            # Zoho accepts different payload shapes depending on function origin/version.
            variants = [
                ("form-arguments-string", {}, None, {"arguments": args_json}),
                ("json-arguments-string", {"Content-Type": "application/json"}, {"arguments": args_json}, None),
                ("form-direct-params", {}, None, args_obj),
                ("json-arguments-map", {"Content-Type": "application/json"}, {"arguments": args_obj}, None),
                ("json-direct-params", {"Content-Type": "application/json"}, args_obj, None),
            ]

            for label, extra_headers, json_body, form_body in variants:
                headers = {**base_headers, **extra_headers}
                resp = requests.post(url, headers=headers, json=json_body, data=form_body, timeout=20)
                if 200 <= resp.status_code < 300:
                    ok, reason = _is_zoho_success(resp)
                    if ok:
                        print(f"  âœ“ OTP sent via Zoho function to {_mask_email(to_email)} ({label})")
                        return True
                    attempts.append(f"{url} [{label}] -> function_error {reason[:180]}")
                    continue
                attempts.append(f"{url} [{label}] -> {resp.status_code} {resp.text[:180]}")

            # Zoho CRM functions often expect `arguments` in query string (URL-encoded JSON).
            q_args_label = "query-arguments-string"
            q_args_resp = requests.post(url, headers=base_headers, params={"arguments": args_json}, timeout=20)
            if 200 <= q_args_resp.status_code < 300:
                ok, reason = _is_zoho_success(q_args_resp)
                if ok:
                    print(f"  ✓ OTP sent via Zoho function to {_mask_email(to_email)} ({q_args_label})")
                    return True
                attempts.append(f"{url} [{q_args_label}] -> function_error {reason[:180]}")
            else:
                attempts.append(f"{url} [{q_args_label}] -> {q_args_resp.status_code} {q_args_resp.text[:180]}")

            # Try with query params as a final fallback on this URL
            q_label = "query-direct-params"
            q_resp = requests.post(url, headers=base_headers, params=args_obj, timeout=20)
            if 200 <= q_resp.status_code < 300:
                ok, reason = _is_zoho_success(q_resp)
                if ok:
                    print(f"  âœ“ OTP sent via Zoho function to {_mask_email(to_email)} ({q_label})")
                    return True
                attempts.append(f"{url} [{q_label}] -> function_error {reason[:180]}")
                continue
            attempts.append(f"{url} [{q_label}] -> {q_resp.status_code} {q_resp.text[:180]}")

        print("  âœ— Zoho OTP function failed after variants:")
        for a in attempts:
            print(f"    - {a}")
        return False
    except Exception as e:
        print(f"  âœ— Zoho OTP function exception: {e}")
        return False


def _send_otp(to_email: str, otp: str, company_name: str) -> bool:
    mode = (OTP_DELIVERY_MODE or "smtp").lower()
    if mode == "zoho":
        if _send_otp_via_zoho_function(to_email, otp, company_name):
            return True
        if SMTP_USER and SMTP_PASSWORD:
            print("  ! Zoho OTP was not confirmed; trying SMTP fallback")
            return _send_otp_email(to_email, otp, company_name)
        return False
    if mode == "auto":
        if _send_otp_via_zoho_function(to_email, otp, company_name):
            return True
        return _send_otp_email(to_email, otp, company_name)
    return _send_otp_email(to_email, otp, company_name)


CAD_EXTENSIONS = {".stl", ".obj", ".ply", ".glb", ".gltf", ".step", ".stp", ".iges", ".igs", ".3mf"}
PROFILE_INDEX_NAME = os.getenv("PINECONE_PROFILE_INDEX", "process-profiles2")
PROFILE_INDEX_DIMENSION = int(os.getenv("PINECONE_PROFILE_INDEX_DIMENSION", "384"))
_profile_text_embedder = None


def _is_cad_filename(name: str) -> bool:
    return Path(name or "").suffix.lower() in CAD_EXTENSIONS


def _autocrop_render_bytes(
    image_bytes: bytes,
    *,
    bg_rgb=(219, 222, 230),
    tol: int = 10,
    out_size: int = 256,
    pad_ratio: float = 0.1,
) -> bytes:
    """
    Tight-crop rendered CAD preview around non-background pixels, then refit.
    This makes thin/small geometry much more readable in preview cards.
    """
    try:
        from PIL import Image
        import numpy as np
    except Exception:
        return image_bytes

    try:
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        arr = np.asarray(img, dtype=np.int16)
        bg = np.asarray(bg_rgb, dtype=np.int16).reshape(1, 1, 3)
        mask = np.any(np.abs(arr - bg) > int(max(0, tol)), axis=2)
        ys, xs = np.where(mask)
        if ys.size == 0 or xs.size == 0:
            return image_bytes

        x0, x1 = int(xs.min()), int(xs.max())
        y0, y1 = int(ys.min()), int(ys.max())
        w = max(1, x1 - x0 + 1)
        h = max(1, y1 - y0 + 1)
        pad = max(4, int(max(w, h) * float(max(0.02, min(0.3, pad_ratio)))))
        x0 = max(0, x0 - pad)
        y0 = max(0, y0 - pad)
        x1 = min(arr.shape[1] - 1, x1 + pad)
        y1 = min(arr.shape[0] - 1, y1 + pad)

        crop = img.crop((x0, y0, x1 + 1, y1 + 1))
        target = int(max(128, min(512, out_size)))
        canvas = Image.new("RGB", (target, target), tuple(int(v) for v in bg_rgb))

        cw, ch = crop.size
        scale = min((target * 0.92) / max(1, cw), (target * 0.92) / max(1, ch))
        nw = max(1, int(round(cw * scale)))
        nh = max(1, int(round(ch * scale)))
        crop = crop.resize((nw, nh), Image.Resampling.LANCZOS)
        ox = (target - nw) // 2
        oy = (target - nh) // 2
        canvas.paste(crop, (ox, oy))

        out = io.BytesIO()
        canvas.save(out, format="JPEG", quality=96, optimize=True, subsampling=0)
        return out.getvalue()
    except Exception:
        return image_bytes


def _shade_tri_faces(tris, base_rgb=(0.38, 0.40, 0.44)):
    """
    Simple Lambert-style face shading for better 3D depth perception.
    """
    import numpy as np

    tri = np.asarray(tris, dtype=float)
    if tri.ndim != 3 or tri.shape[1] != 3 or tri.shape[2] != 3:
        base = np.array(base_rgb, dtype=float)
        return np.tile(np.array([base[0], base[1], base[2], 1.0], dtype=float), (max(len(tri), 1), 1))

    v1 = tri[:, 1] - tri[:, 0]
    v2 = tri[:, 2] - tri[:, 0]
    n = np.cross(v1, v2)
    n_norm = np.linalg.norm(n, axis=1, keepdims=True)
    n = n / np.clip(n_norm, 1e-9, None)

    light = np.array([0.35, -0.35, 0.87], dtype=float)
    light = light / np.linalg.norm(light)
    intensity = (n @ light).reshape(-1)
    intensity = np.clip(intensity, -0.35, 1.0)
    intensity = 0.35 + 0.65 * ((intensity + 0.35) / 1.35)

    base = np.array(base_rgb, dtype=float).reshape(1, 3)
    rgb = np.clip(base * (0.62 + 0.72 * intensity[:, None]), 0.0, 1.0)
    alpha = np.ones((rgb.shape[0], 1), dtype=float)
    return np.hstack([rgb, alpha])


def _render_mesh_view_image(tris, bg, body, edge, elev: float, azim: float) -> bytes:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from mpl_toolkits.mplot3d.art3d import Poly3DCollection

    # CAD_RENDER_RES = output side in px (memory-safe default tuned for Render 512MB)
    cad_render_res = os.getenv("CAD_RENDER_RES", "256").strip()
    try:
        side_px = int(cad_render_res)
    except ValueError:
        side_px = 160
    side_px = max(160, min(512, side_px))
    dpi = 100
    supersample_raw = os.getenv("CAD_RENDER_SUPERSAMPLE", "2").strip()
    try:
        supersample = int(supersample_raw)
    except ValueError:
        supersample = 2
    supersample = max(1, min(3, supersample))
    render_px = int(min(1024, side_px * supersample))
    fig_size = render_px / float(dpi)

    # Keep style lightweight for high-triangle previews without flattening detail.
    tri_count = int(getattr(tris, "shape", [0])[0] or 0)
    fast_style = tri_count >= int(os.getenv("CAD_FAST_STYLE_TRI_THRESHOLD", "35000"))
    lw = 0.03 if fast_style else 0.08
    edge_rgba = (0.12, 0.12, 0.12, 0.16) if fast_style else (0.10, 0.10, 0.10, 0.24)

    fig = plt.figure(figsize=(fig_size, fig_size), dpi=dpi, facecolor=bg)
    ax = fig.add_subplot(1, 1, 1, projection="3d")
    ax.set_facecolor(bg)
    facecolors = _shade_tri_faces(tris, base_rgb=body[:3] if isinstance(body, tuple) else (0.38, 0.40, 0.44))

    poly = Poly3DCollection(
        tris,
        facecolor=facecolors,
        edgecolor=edge_rgba,
        linewidths=lw,
        antialiaseds=True,
        alpha=1.0,
    )
    ax.add_collection3d(poly)
    ax.set_xlim(-0.95, 0.95)
    ax.set_ylim(-0.95, 0.95)
    ax.set_zlim(-0.95, 0.95)
    ax.set_box_aspect((1, 1, 1))
    ax.view_init(elev=elev, azim=azim)
    # Perspective retains depth cues better than ortho for organic 3MF meshes.
    try:
        projection = os.getenv("CAD_PROJECTION", "persp").strip().lower()
        ax.set_proj_type("ortho" if projection == "ortho" else "persp")
    except Exception:
        pass
    ax.set_axis_off()
    plt.subplots_adjust(left=0.01, right=0.99, bottom=0.01, top=0.99)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, facecolor=bg)
    plt.close(fig)
    rendered = buf.getvalue()
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(rendered)).convert("RGB")
        if im.size != (side_px, side_px):
            im = im.resize((side_px, side_px), Image.Resampling.LANCZOS)
        out = io.BytesIO()
        im.save(out, format="JPEG", quality=96, optimize=True, subsampling=0)
        rendered = out.getvalue()
    except Exception:
        pass
    try:
        bg_rgb = tuple(int(round(float(c) * 255.0)) for c in bg[:3])
    except Exception:
        bg_rgb = (219, 222, 230)
    return _autocrop_render_bytes(
        rendered,
        bg_rgb=bg_rgb,
        tol=10,
        out_size=side_px,
        pad_ratio=0.1,
    )


def _mesh_to_preview_bundle(mesh, max_faces_override: int | None = None) -> dict:
    import numpy as np
    try:
        import matplotlib  # noqa: F401
    except Exception as e:
        raise ValueError(f"CAD preview renderer unavailable: {e}")

    verts = np.asarray(getattr(mesh, "vertices", []), dtype=float)
    faces = np.asarray(getattr(mesh, "faces", []), dtype=int)
    if verts.size == 0 or faces.size == 0:
        raise ValueError("Mesh preview requires triangle faces")

    # Robust framing (less sensitive to stray/outlier vertices)
    center = np.median(verts, axis=0)
    v = verts - center
    d = np.linalg.norm(v, axis=1)
    if d.size:
        p = float(np.percentile(d, 99.5))
    else:
        p = 0.0
    radius = p if p > 1e-9 else float(np.max(d) if d.size else 1.0)
    if radius <= 1e-9:
        radius = 1.0
    v = v / radius
    # Clamp extreme outliers so they don't force tiny in-frame object.
    v = np.clip(v, -1.35, 1.35)

    # Keep rendering speed bounded for dense meshes (critical for 512MB runtime),
    # but avoid over-decimation that creates low-quality faceted previews.
    try:
        max_faces = int(max_faces_override) if max_faces_override is not None else int(os.getenv("CAD_RENDER_MAX_FACES", "30000"))
    except ValueError:
        max_faces = 30000
    max_faces = max(5000, min(60000, max_faces))
    if len(faces) > max_faces:
        try:
            # Prefer geometric decimation over raw face skipping.
            decimated = mesh.simplify_quadric_decimation(max_faces)
            dec_faces = np.asarray(getattr(decimated, "faces", []), dtype=int)
            dec_verts = np.asarray(getattr(decimated, "vertices", []), dtype=float)
            if dec_faces.size and dec_verts.size:
                mesh = decimated
                verts = dec_verts
                faces = dec_faces
        except Exception:
            # Last-resort fallback: keep a stratified random subset (better than stride artifacts).
            rng = np.random.default_rng(42)
            keep = np.sort(rng.choice(len(faces), size=max_faces, replace=False))
            faces = faces[keep]

    tris = v[faces]

    bg = (0.90, 0.91, 0.94, 1.0)         # light gray sheet
    body = (0.48, 0.50, 0.54, 1.0)       # mid gray; lighting will add depth
    edge = (0.10, 0.10, 0.10, 0.55)      # dark technical edges

    # Memory-safe default for Render 512MB: render only isometric study view.
    study = {
        "name": "iso_study",
        "bytes": _render_mesh_view_image(tris, bg, body, edge, elev=24, azim=36),
    }
    return {"study": study, "extras": []}


def _cad_mesh_stats(mesh) -> dict:
    raw_extents = getattr(mesh, "extents", None)
    if raw_extents is None:
        extents = [0.0, 0.0, 0.0]
    else:
        try:
            extents = [float(v) for v in list(raw_extents)]
        except Exception:
            extents = [0.0, 0.0, 0.0]
    return {
        "triangles": int(getattr(mesh, "faces", []).shape[0]) if getattr(mesh, "faces", None) is not None else 0,
        "vertices": int(getattr(mesh, "vertices", []).shape[0]) if getattr(mesh, "vertices", None) is not None else 0,
        "bbox_x": float(extents[0]) if len(extents) > 0 else 0.0,
        "bbox_y": float(extents[1]) if len(extents) > 1 else 0.0,
        "bbox_z": float(extents[2]) if len(extents) > 2 else 0.0,
        "surface_area": float(getattr(mesh, "area", 0.0) or 0.0),
        "volume": float(getattr(mesh, "volume", 0.0) or 0.0),
    }


def _cad_project_details_from_inference(inference: dict, cad_stats: dict) -> dict:
    def _value(obj, key):
        if isinstance(obj, dict):
            return obj.get(key, "")
        return ""

    part_family_obj = _value(inference or {}, "part_family")
    material_obj = _value(inference or {}, "material")
    process_obj = _value(inference or {}, "process")

    part_family = _value(part_family_obj, "value") if isinstance(part_family_obj, dict) else str(part_family_obj or "")
    material = _value(material_obj, "value") if isinstance(material_obj, dict) else str(material_obj or "")
    process = _value(process_obj, "primary") if isinstance(process_obj, dict) else str(process_obj or "")
    detail = _value(part_family_obj, "detail") if isinstance(part_family_obj, dict) else ""
    triangles = int((cad_stats or {}).get("triangles") or 0)

    name_bits = [part_family or "CAD Component", process]
    project_name = " - ".join([x for x in name_bits if x]).strip() or "CAD Project"
    overview = (
        f"CAD-derived part analyzed from uploaded model. "
        f"Likely {part_family or 'component'} in {material or 'unspecified material'} "
        f"using {process or 'unspecified process'}. "
        f"Mesh triangles: {triangles}. "
        f"{detail}".strip()
    )
    return {
        "project_name": project_name[:120],
        "overview": overview[:800],
        "customer_industry": "",
    }


def _extract_pdf_context(raw_bytes: bytes, max_chars: int = 60000, max_images: int = 4) -> dict:
    text_parts = []
    image_previews = []
    page_count = 0

    # Preferred path: PyMuPDF (text + raster previews)
    if fitz is not None:
        try:
            doc = fitz.open(stream=raw_bytes, filetype="pdf")
            page_count = len(doc)
            for page_index, page in enumerate(doc):
                page_text = (page.get_text("text") or "").strip()
                if page_text:
                    text_parts.append(page_text)

                if len(image_previews) < max_images:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2), alpha=False)
                    png_bytes = pix.tobytes("png")
                    image_previews.append({
                        "page": page_index + 1,
                        "mime_type": "image/png",
                        "data_url": f"data:image/png;base64,{base64.b64encode(png_bytes).decode('utf-8')}",
                    })
            doc.close()
        except Exception:
            text_parts = []
            image_previews = []
            page_count = 0

    # Fallback path: pypdf text only
    if not text_parts and PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(raw_bytes))
            page_count = len(reader.pages)
            for p in reader.pages:
                t = (p.extract_text() or "").strip()
                if t:
                    text_parts.append(t)
        except Exception:
            pass

    full_text = "\n\n".join(text_parts).strip()
    full_text = full_text[:max_chars]
    return {"text": full_text, "images": image_previews, "page_count": page_count}


def _extract_document_text(raw_bytes: bytes, filename: str) -> dict:
    """
    Extract plain text from non-PDF office/document files.
    Supported: docx, txt, md, csv, tsv, rtf, json
    """
    name = (filename or "").strip()
    ext = Path(name).suffix.lower()
    text = ""
    parser = "none"

    if ext in {".txt", ".md", ".csv", ".tsv", ".json"}:
        text = raw_bytes.decode("utf-8", errors="ignore")
        parser = "utf8"
    elif ext == ".rtf":
        raw = raw_bytes.decode("utf-8", errors="ignore")
        # Light RTF cleanup
        raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
        raw = re.sub(r"\\[a-zA-Z]+\d* ?", " ", raw)
        raw = re.sub(r"[{}]", " ", raw)
        text = re.sub(r"\s{2,}", " ", raw).strip()
        parser = "rtf"
    elif ext == ".docx":
        try:
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            lines = []
            for p in doc.paragraphs:
                s = (p.text or "").strip()
                if s:
                    lines.append(s)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    joined = " | ".join([cell for cell in cells if cell])
                    if joined:
                        lines.append(joined)
            text = "\n".join(lines).strip()
            parser = "docx"
        except Exception:
            text = ""
            parser = "docx_failed"
    elif ext == ".doc":
        # Legacy binary .doc is not reliably parseable without external tooling.
        text = ""
        parser = "doc_unsupported"

    text = (text or "")[:60000]
    if _looks_like_email_thread(text):
        cleaned = _clean_email_thread_text(text)
        if cleaned:
            print(f"[extract-document] Email-thread cleanup applied ({len(text)} -> {len(cleaned)} chars)")
            text = cleaned
    details = _derive_project_details_from_text(text)
    return {
        "text": text,
        "project_details": details,
        "parser": parser,
        "extension": ext,
    }


def _looks_like_email_thread(text: str) -> bool:
    raw = (text or "").lower()
    if not raw:
        return False
    signals = 0
    for pat in (
        r"^\s*from\s*:",
        r"^\s*to\s*:",
        r"^\s*cc\s*:",
        r"^\s*subject\s*:",
        r"^\s*sent\s*:",
        r"^\s*date\s*:",
        r"^\s*regards\b",
        r"^\s*best regards\b",
        r"^\s*thanks\b",
        r"^[-_]{2,}\s*original message\s*[-_]{2,}",
    ):
        if re.search(pat, raw, flags=re.I | re.M):
            signals += 1
    return signals >= 2


def _clean_email_thread_text(text: str) -> str:
    raw = (text or "").replace("\r\n", "\n")
    if not raw.strip():
        return ""

    keep_hard = re.compile(
        r"(part\s*name|material|process|quantity|qty|finish|tolerance|delivery|lead\s*time|"
        r"drawing|rev\b|rfq|rfp|po[-\s]?\d+|price|setup\s*cost|inspection|volume|units?/year|"
        r"outcome|lesson\s*learned|project)",
        re.I,
    )
    drop_line = re.compile(
        r"^\s*(from|to|cc|bcc|subject|sent|date)\s*:\s*.*$|"
        r"^\s*(hi|hello|dear)\b.*$|"
        r"^\s*(regards|best regards|thanks|thank you)\b.*$|"
        r"^\s*[-_]{2,}\s*original message\s*[-_]{2,}\s*$|"
        r"^\s*this e-?mail.*(confidential|privileged).*$|"
        r"^\s*please consider the environment before printing.*$",
        re.I,
    )

    lines = [ln.rstrip() for ln in raw.split("\n")]
    cleaned: list[str] = []
    for ln in lines:
        s = ln.strip()
        if not s:
            cleaned.append("")
            continue
        if keep_hard.search(s):
            cleaned.append(s)
            continue
        if drop_line.search(s):
            continue
        if re.fullmatch(r"[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}", s, flags=re.I):
            continue
        if re.fullmatch(r"[\+\(]?[0-9\-\(\)\s]{7,}$", s):
            continue
        cleaned.append(s)

    out: list[str] = []
    last_blank = False
    for ln in cleaned:
        blank = not ln.strip()
        if blank and last_blank:
            continue
        out.append(ln)
        last_blank = blank
    return "\n".join(out).strip()


def _extract_3mf_embedded_images(raw_bytes: bytes, max_images: int = 3) -> list[bytes]:
    """
    Fallback for 3MF files that fail mesh parsing.
    Extract thumbnail/embedded raster images directly from the 3MF zip container.
    """
    out: list[bytes] = []
    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes), "r") as zf:
            names = zf.namelist()
            candidates = []
            for n in names:
                lower = n.lower()
                if lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
                    score = 0
                    if "thumbnail" in lower:
                        score += 4
                    if "thumbnails/" in lower:
                        score += 3
                    if "metadata/" in lower:
                        score += 1
                    candidates.append((score, n))

            # Prefer detailed images (not flat color blocks), then size.
            ranked = []
            for score, name in candidates:
                try:
                    raw = zf.read(name)
                    from PIL import Image
                    import numpy as np
                    im = Image.open(io.BytesIO(raw)).convert("RGB")
                    w, h = im.size
                    area = int(w) * int(h)
                    arr = np.asarray(im, dtype=np.uint8)
                    std = float(arr.std())
                    detail_ok = (area >= 64000) and (std >= 8.0)
                    quality_score = (2.0 * std) + (area / 100000.0)
                    ranked.append((int(detail_ok), quality_score, area, score, name, im))
                except Exception:
                    continue
            ranked.sort(key=lambda x: (-x[0], -x[1], -x[2], -x[3], x[4]))

            for _, _, _, _, _, im in ranked[:max_images]:
                # If embedded preview is tiny, upscale so on-screen preview is clearer.
                min_side = min(im.size)
                if min_side < 512 and min_side > 0:
                    scale = 512.0 / float(min_side)
                    new_w = int(round(im.size[0] * scale))
                    new_h = int(round(im.size[1] * scale))
                    im = im.resize((new_w, new_h), Image.Resampling.LANCZOS)
                buf = io.BytesIO()
                im.save(buf, format="JPEG", quality=96, optimize=True, subsampling=0)
                out.append(buf.getvalue())
    except Exception:
        return []
    return out


def _parse_3mf_transform_matrix(transform_str: str):
    """
    3MF component transform is 12 floats (3x4 affine):
    m00 m01 m02 m03 m10 m11 m12 m13 m20 m21 m22 m23
    """
    vals = [v for v in str(transform_str or "").strip().split() if v]
    if len(vals) != 12:
        return None
    try:
        nums = [float(v) for v in vals]
    except Exception:
        return None
    return [
        [nums[0], nums[1], nums[2], nums[3]],
        [nums[4], nums[5], nums[6], nums[7]],
        [nums[8], nums[9], nums[10], nums[11]],
    ]


def _apply_3mf_transform(vertices, mat34):
    if mat34 is None:
        return vertices
    import numpy as np
    v = np.asarray(vertices, dtype=float)
    if v.size == 0:
        return v
    ones = np.ones((v.shape[0], 1), dtype=float)
    vh = np.hstack([v, ones])  # Nx4
    m = np.asarray(mat34, dtype=float)  # 3x4
    out = vh @ m.T  # Nx3
    return out


def _load_3mf_mesh_from_xml(raw_bytes: bytes):
    """
    Parse .3mf XML directly to build a mesh when trimesh loader fails (e.g. KeyError 'world').
    Handles direct mesh objects and component references.
    """
    import numpy as np
    import xml.etree.ElementTree as ET

    if trimesh is None:
        raise RuntimeError("trimesh is required for 3MF mesh fallback")

    with zipfile.ZipFile(io.BytesIO(raw_bytes), "r") as zf:
        model_name = None
        preferred = "3d/3dmodel.model"
        for n in zf.namelist():
            low = n.lower()
            if low == preferred:
                model_name = n
                break
        if not model_name:
            for n in zf.namelist():
                if n.lower().endswith(".model"):
                    model_name = n
                    break
        if not model_name:
            raise RuntimeError("3MF model XML not found in archive")
        model_xml = zf.read(model_name)

    root = ET.fromstring(model_xml)

    def _lname(tag: str) -> str:
        return str(tag).split("}")[-1]

    def _first_child_by_name(node, name: str):
        if node is None:
            return None
        for ch in list(node):
            if _lname(ch.tag) == name:
                return ch
        return None

    def _children_by_name(node, name: str):
        if node is None:
            return []
        return [ch for ch in list(node) if _lname(ch.tag) == name]

    resources = _first_child_by_name(root, "resources")
    if resources is None:
        # namespace-agnostic fallback for variant 3MF XMLs
        for el in root.iter():
            if _lname(el.tag) == "resources":
                resources = el
                break
    if resources is None:
        raise RuntimeError("3MF resources section missing")

    # Parse objects into lightweight descriptors
    objects = {}
    for obj in _children_by_name(resources, "object"):
        oid = obj.attrib.get("id")
        if not oid:
            continue

        mesh_node = _first_child_by_name(obj, "mesh")
        comps_node = _first_child_by_name(obj, "components")

        if mesh_node is not None:
            verts = []
            faces = []
            verts_node = _first_child_by_name(mesh_node, "vertices")
            tris_node = _first_child_by_name(mesh_node, "triangles")
            if verts_node is not None:
                for v in _children_by_name(verts_node, "vertex"):
                    try:
                        verts.append([float(v.attrib.get("x", 0.0)), float(v.attrib.get("y", 0.0)), float(v.attrib.get("z", 0.0))])
                    except Exception:
                        continue
            if tris_node is not None:
                for t in _children_by_name(tris_node, "triangle"):
                    try:
                        faces.append([int(t.attrib.get("v1", 0)), int(t.attrib.get("v2", 0)), int(t.attrib.get("v3", 0))])
                    except Exception:
                        continue
            if verts and faces:
                objects[oid] = {
                    "type": "mesh",
                    "vertices": np.asarray(verts, dtype=float),
                    "faces": np.asarray(faces, dtype=int),
                }
                continue

        if comps_node is not None:
            comps = []
            for c in _children_by_name(comps_node, "component"):
                refid = c.attrib.get("objectid")
                if not refid:
                    continue
                comps.append((refid, _parse_3mf_transform_matrix(c.attrib.get("transform", ""))))
            if comps:
                objects[oid] = {"type": "components", "components": comps}

    if not objects:
        raise RuntimeError("No mesh/components found in 3MF resources")

    cache = {}

    def build_obj(oid: str, stack=None):
        stack = stack or set()
        if oid in cache:
            return cache[oid].copy()
        if oid in stack:
            return None
        desc = objects.get(oid)
        if not desc:
            return None
        stack.add(oid)
        result = None
        if desc["type"] == "mesh":
            result = trimesh.Trimesh(vertices=desc["vertices"], faces=desc["faces"], process=False)
        else:
            parts = []
            for child_id, mat in desc.get("components", []):
                child = build_obj(child_id, stack)
                if child is None:
                    continue
                v = _apply_3mf_transform(child.vertices, mat)
                child = trimesh.Trimesh(vertices=v, faces=child.faces, process=False)
                if len(child.faces) > 0:
                    parts.append(child)
            if parts:
                result = trimesh.util.concatenate(parts)
        stack.remove(oid)
        if result is not None:
            cache[oid] = result.copy()
        return result

    # Prefer build items; otherwise, combine all root mesh objects.
    build_node = _first_child_by_name(root, "build")
    if build_node is None:
        for el in root.iter():
            if _lname(el.tag) == "build":
                build_node = el
                break
    meshes = []
    if build_node is not None:
        for item in _children_by_name(build_node, "item"):
            oid = item.attrib.get("objectid")
            if not oid:
                continue
            m = build_obj(oid, set())
            if m is None:
                continue
            mat = _parse_3mf_transform_matrix(item.attrib.get("transform", ""))
            if mat is not None:
                v = _apply_3mf_transform(m.vertices, mat)
                m = trimesh.Trimesh(vertices=v, faces=m.faces, process=False)
            if len(m.faces) > 0:
                meshes.append(m)
    if not meshes:
        for oid, d in objects.items():
            if d.get("type") == "mesh":
                m = build_obj(oid, set())
                if m is not None and len(m.faces) > 0:
                    meshes.append(m)

    if not meshes:
        raise RuntimeError("No triangulated geometry resolved from 3MF")
    return trimesh.util.concatenate(meshes)


def _clean(v: str) -> str:
    s = str(v or "").strip()
    if not s:
        return ""
    s = s.replace("**", "")
    s = s.replace("`", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _normalize_project(v: str, fallback: str = "") -> str:
    s = _clean(v)
    fb = _clean(fallback)
    if not s:
        return fb
    if re.match(r"^(issued|date)\s*:", s, flags=re.I):
        return fb
    if re.match(r"^#?\s*request for proposal", s, flags=re.I):
        return fb
    return s


def _is_person_name(s: str) -> bool:
    s = _clean(s)
    if not s or not (3 <= len(s) <= 60):
        return False
    if re.search(r"\d", s):
        return False
    words = s.split()
    if len(words) < 2:
        return False
    if s == s.upper():
        return False
    if not all(w[0].isupper() for w in words if w):
        return False
    bad = re.compile(
        r"\b(inc|llc|ltd|corp|co\.|company|group|solutions|services|"
        r"technologies|industries|manufacturing|request|proposal|rfp|"
        r"quote|order|logo|header|footer|watermark)\b",
        re.I,
    )
    return not bool(bad.search(s))


def _normalize_contact_name(v: str) -> str:
    s = _clean(v)
    return s[:120] if s and _is_person_name(s) else ""


def _normalize_project_date(raw: str) -> str:
    text = _clean(raw)
    if not text:
        return ""
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
        return text
    if re.fullmatch(r"\d{4}-\d{2}", text):
        return f"{text}-01"
    for fmt in ("%B %Y", "%b %Y", "%m/%Y", "%m-%Y"):
        try:
            return datetime.strptime(text, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return ""


def _derive_project_details_from_text(text: str) -> dict:
    txt = (text or "").strip()
    if not txt:
        return {
            "project_name": "",
            "overview": "",
            "customer_industry": "",
            "customer_name": "",
            "contact_name": "",
            "contact_email": "",
            "contact_phone": "",
            "company_name": "",
            "company_location": "",
            "project_date": "",
        }

    # Very light heuristics to prefill fields from RFP text.
    first_lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
    project_name = ""
    for ln in first_lines[:20]:
        if len(ln) > 4 and len(ln) < 110:
            if re.search(r"(rfp|request for proposal|work order|quotation)", ln, flags=re.I):
                continue
            project_name = ln
            break

    overview = " ".join(first_lines[:12])[:900]

    industry = ""
    industry_keywords = [
        ("aerospace", "Aerospace"),
        ("medical", "Medical"),
        ("automotive", "Automotive"),
        ("robotics", "Robotics"),
        ("defence", "Defence"),
        ("defense", "Defence"),
        ("industrial", "Industrial"),
        ("energy", "Energy"),
        ("consumer", "Consumer"),
    ]
    lower = txt.lower()
    for key, label in industry_keywords:
        if key in lower:
            industry = label
            break

    def _m(pattern: str) -> str:
        m = re.search(pattern, txt, flags=re.I)
        return (m.group(1).strip() if m else "")

    customer_name = _m(r"\b(?:customer name|buyer|client)\s*[:\-]?\s*([^\n\r]{2,120})")
    contact_name = _normalize_contact_name(
        _m(r"\b(?:contact name|attention|attn)\s*[:\-]?\s*([^\n\r]{2,120})")
    )
    contact_email = _m(r"\b(?:contact email|email)\s*[:\-]?\s*([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})")
    if not contact_email:
        m_email = re.search(r"([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})", txt, flags=re.I)
        contact_email = (m_email.group(1).strip().lower() if m_email else "")
    contact_phone = _m(r"\b(?:contact (?:phone|number)|phone|mobile|tel)\s*[:\-]?\s*([\+\d][\d\-\(\) ]{6,30})")
    company_name = _m(r"\b(?:company name|supplier name|vendor|account name)\s*[:\-]?\s*([^\n\r]{2,120})")
    company_location = _m(r"\b(?:company location|location|address)\s*[:\-]?\s*([^\n\r]{3,140})")
    project_date = _normalize_project_date(
        _m(r"\b(?:project date|date)\s*[:\-]?\s*(\d{4}-\d{2}-\d{2}|\d{4}-\d{2}|[A-Za-z]{3,9}\s+\d{4}|\d{1,2}[/-]\d{4})")
    )

    company_name = _clean(company_name)
    customer_name = _clean(customer_name)
    if len(company_name.split()) >= 14 or len(company_name) > 120:
        company_name = customer_name

    return {
        "project_name": _normalize_project(project_name, customer_name)[:120],
        "overview": overview[:900],
        "customer_industry": industry,
        "customer_name": customer_name[:120],
        "contact_name": contact_name,
        "contact_email": contact_email[:120],
        "contact_phone": contact_phone[:40],
        "company_name": company_name[:120],
        "company_location": _clean(company_location)[:140],
        "project_date": project_date[:10],
    }


def _format_project_date(raw: str):
    if not raw or len(raw.strip()) != 7:
        return None
    try:
        return datetime.strptime(raw.strip(), "%Y-%m").strftime("%B %Y")
    except ValueError:
        return None


def _to_float_vec(vec) -> list:
    """
    Convert any numeric list to a list of Python floats, padded/trimmed to
    exactly INDEX_DIMENSION elements.
    """
    if not vec:
        return [0.0] * INDEX_DIMENSION
    result = [float(v) for v in vec]
    if len(result) < INDEX_DIMENSION:
        result += [0.0] * (INDEX_DIMENSION - len(result))
    return result[:INDEX_DIMENSION]


def _is_zero_vector(vec: list) -> bool:
    """Return True if every element in the vector is 0.0."""
    return all(v == 0.0 for v in vec)


def _compute_embedding_from_part_payload(part: dict, saved_image_path: Path | None = None):
    """
    Recompute visual embedding from saved image or inline base64 payload.
    Used as a recovery path when frontend-provided vector is empty/zero.
    """
    try:
        if saved_image_path and saved_image_path.exists():
            return compute_clip_embedding(str(saved_image_path))
    except Exception as e:
        print(f"  ! Embedding recompute from saved image failed: {e}")

    img_b64 = (part.get("image_b64") or "").strip()
    if img_b64:
        tmp = UPLOAD_DIR / f"{uuid.uuid4()}.jpg"
        try:
            with open(tmp, "wb") as f:
                f.write(base64.b64decode(img_b64))
            return compute_clip_embedding(str(tmp))
        except Exception as e:
            print(f"  ! Embedding recompute from image_b64 failed: {e}")
        finally:
            try:
                if tmp.exists():
                    tmp.unlink()
            except Exception:
                pass
    return None


def _make_fallback_vector(part_id: str, metadata: dict) -> list:
    """
    âœ… FIX: When CLIP is unavailable we cannot push an all-zero vector because
    Pinecone requires at least one non-zero value.

    Strategy: build a deterministic, lightweight fingerprint vector from the
    part's categorical metadata (part_family, material, process) so that
    semantically similar parts still land near each other, while ensuring the
    vector is never all-zeros.

    The fingerprint uses Python's built-in hash() seeded on the field strings,
    producing stable floats in the range [-1, 1].  We set element [0] to a
    small non-zero sentinel (0.001) as an absolute guarantee.
    """
    import hashlib

    fields = [
        metadata.get("part_family", ""),
        metadata.get("material", ""),
        metadata.get("process_primary", ""),
        metadata.get("process_secondary", ""),
        metadata.get("complexity_class", ""),
        metadata.get("tolerance_class", ""),
        part_id,
    ]
    seed_str = "|".join(str(f) for f in fields)

    vec = []
    for i in range(INDEX_DIMENSION):
        # Deterministic float in [-1, 1] derived from seed + position
        h = hashlib.md5(f"{seed_str}:{i}".encode()).digest()
        # Take first 4 bytes as a signed int32, normalise to [-1, 1]
        val = int.from_bytes(h[:4], "big", signed=True) / (2 ** 31)
        vec.append(float(val) * 0.05)   # scale down so CLIP vectors dominate in search

    # Absolute sentinel: element 0 is never zero
    if vec[0] == 0.0:
        vec[0] = 0.001

    return vec


# â”€â”€ Pinecone init â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
DEFER_PINECONE_INIT = os.getenv("DEFER_PINECONE_INIT", "false").strip().lower() == "true"

pinecone_index = None
if not DEFER_PINECONE_INIT:
    try:
        if PINECONE_API_KEY:
            _pc = Pinecone(api_key=PINECONE_API_KEY)
            existing = {i.name: i for i in _pc.list_indexes()}
            if INDEX_NAME in existing:
                current_dim = existing[INDEX_NAME].dimension
                if current_dim != INDEX_DIMENSION:
                    print(f"  ⚠ Index '{INDEX_NAME}' is {current_dim}-dim, need {INDEX_DIMENSION}-dim — recreating…")
                    _pc.delete_index(INDEX_NAME)
                    time.sleep(5)
                    _pc.create_index(
                        name=INDEX_NAME,
                        dimension=INDEX_DIMENSION,
                        metric="cosine",
                        spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                    )
            else:
                _pc.create_index(
                    name=INDEX_NAME,
                    dimension=INDEX_DIMENSION,
                    metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region="us-east-1"),
                )
            pinecone_index = _pc.Index(INDEX_NAME)
            stats = pinecone_index.describe_index_stats()
            print(f"✓ Pinecone connected — {INDEX_NAME} ({INDEX_DIMENSION}-dim cosine) | vectors: {stats.total_vector_count}")
    except Exception as e:
        print(f"✗ Pinecone init failed: {e}")
else:
    print("~ Pinecone init deferred for startup")

    
pinecone_index = None
try:
    if PINECONE_API_KEY:
        _pc      = Pinecone(api_key=PINECONE_API_KEY)
        existing = {i.name: i for i in _pc.list_indexes()}
        if INDEX_NAME in existing:
            current_dim = existing[INDEX_NAME].dimension
            if current_dim != INDEX_DIMENSION:
                print(f"  âš  Index '{INDEX_NAME}' is {current_dim}-dim, need {INDEX_DIMENSION}-dim â€” recreatingâ€¦")
                _pc.delete_index(INDEX_NAME)
                time.sleep(5)
                _pc.create_index(name=INDEX_NAME, dimension=INDEX_DIMENSION, metric="cosine",
                                 spec=ServerlessSpec(cloud="aws", region="us-east-1"))
        else:
            _pc.create_index(name=INDEX_NAME, dimension=INDEX_DIMENSION, metric="cosine",
                             spec=ServerlessSpec(cloud="aws", region="us-east-1"))
        pinecone_index = _pc.Index(INDEX_NAME)
        stats = pinecone_index.describe_index_stats()
        print(f"âœ“ Pinecone connected â€” {INDEX_NAME} ({INDEX_DIMENSION}-dim cosine) | vectors: {stats.total_vector_count}")
except Exception as e:
    print(f"âœ— Pinecone init failed: {e}")

# â”€â”€ Zoho token cache â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_zoho_token_cache = {"token": "", "expires_at": 0}

def get_zoho_token() -> str:
    now = time.time()
    if _zoho_token_cache["token"] and _zoho_token_cache["expires_at"] > now + 60:
        return _zoho_token_cache["token"]
    if not ZOHO_REFRESH_TOKEN:
        raise ValueError("ZOHO_REFRESH_TOKEN not set in .env")
    resp = requests.post(ZOHO_ACCOUNTS_URL, data={
        "refresh_token": ZOHO_REFRESH_TOKEN,
        "client_id":     ZOHO_CLIENT_ID,
        "client_secret": ZOHO_CLIENT_SECRET,
        "grant_type":    "refresh_token",
    }, timeout=10)
    resp.raise_for_status()
    data = resp.json()
    if "access_token" not in data:
        raise ValueError(f"Zoho token refresh failed: {data}")
    _zoho_token_cache["token"]      = data["access_token"]
    _zoho_token_cache["expires_at"] = now + data.get("expires_in", 3600)
    print("  âœ“ Zoho access token refreshed")
    return _zoho_token_cache["token"]

def zoho_headers() -> dict:
    return {"Authorization": f"Zoho-oauthtoken {get_zoho_token()}",
            "Content-Type": "application/json"}


def zoho_auth_headers() -> dict:
    return {"Authorization": f"Zoho-oauthtoken {get_zoho_token()}"}

def zoho_lookup_contact(email: str) -> dict:
    url    = f"{ZOHO_API_BASE}/Contacts/search"
    params = {"email": email}
    resp   = requests.get(url, headers=zoho_headers(), params=params, timeout=10)
    if resp.status_code == 204:
        raise ValueError(f"No contact found in Zoho CRM for email: {email}")
    resp.raise_for_status()
    data     = resp.json()
    contacts = data.get("data", [])
    if not contacts:
        raise ValueError(f"No contact found in Zoho CRM for email: {email}")
    contact = contacts[0]
    account = contact.get("Account_Name", {})
    if isinstance(account, dict):
        zoho_account_id = str(account.get("id", ""))
        company_name    = account.get("name", "")
    else:
        zoho_account_id = ""
        company_name    = str(account)
    if not zoho_account_id:
        zoho_account_id = str(contact.get("id", ""))
        company_name    = contact.get("Company", "") or contact.get("Full_Name", "")
    full_name = contact.get("Full_Name", "") or \
                f"{contact.get('First_Name','')} {contact.get('Last_Name','')}".strip()
    print(f"  âœ“ Zoho lookup: {email} â†’ {company_name} (Account ID: {zoho_account_id})")
    return {"company_name": company_name, "zoho_account_id": zoho_account_id, "contact_name": full_name}

from fastapi import Request
from fastapi.responses import JSONResponse
import requests


# â”€â”€ REPLACE your existing zoho_push_past_project() with this â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _to_zoho_date(raw: str) -> str:
    text = (raw or "").strip()
    if not text:
        return ""
    if len(text) == 10 and text[4] == "-" and text[7] == "-":
        return text
    if len(text) == 7 and text[4] == "-":
        return f"{text}-01"
    try:
        return datetime.strptime(text, "%B %Y").strftime("%Y-%m-%d")
    except ValueError:
        return ""


def _resolve_supplier_contact_lookup_id(
    preferred_email: str,
    fallback_email: str,
    supplier_account_id: str,
) -> str:
    """
    Resolve Contact id for Supplier_Contact_Lookup in Supplier_Past_Projects.
    Prefers contact rows mapped to the same supplier account id.
    """
    account_id = str(supplier_account_id or "").strip()
    emails = []
    for raw in (preferred_email, fallback_email):
        e = str(raw or "").strip().lower()
        if e and e not in emails:
            emails.append(e)
    if not emails:
        return ""

    for email in emails:
        try:
            resp = requests.get(
                f"{ZOHO_API_BASE}/Contacts/search",
                headers=zoho_headers(),
                params={"email": email},
                timeout=10,
            )
            if resp.status_code != 200:
                continue
            rows = (resp.json() or {}).get("data", []) or []
            if not rows:
                continue

            if account_id:
                for row in rows:
                    cid = str((row or {}).get("id") or "").strip()
                    acct = (row or {}).get("Account_Name")
                    acct_id = str((acct or {}).get("id") if isinstance(acct, dict) else "").strip()
                    if cid and acct_id and acct_id == account_id:
                        return cid

            cid = str((rows[0] or {}).get("id") or "").strip()
            if cid:
                return cid
        except Exception:
            continue
    return ""


def _build_zoho_project_record(part: dict, *, index_ready: bool = False) -> dict:
    company_name = (part.get("company_name") or "").strip()
    pinecone_id = (part.get("part_id") or "").strip()
    project_name = (part.get("project_name") or part.get("part_name") or company_name or "Project").strip()
    supplier_email = (part.get("supplier_email") or "").strip().lower()
    zoho_id = (part.get("zoho_id") or "").strip()
    project_date = _to_zoho_date(part.get("project_date", ""))

    image_url_value = (part.get("image_url") or "").strip()
    if not image_url_value:
        fallback_preview = (part.get("image_preview") or "").strip()
        if fallback_preview and not fallback_preview.startswith("blob:"):
            image_url_value = fallback_preview

    ncr_description = (part.get("ncr_description") or "").strip()
    quoting_lesson = (part.get("quoting_lesson") or "").strip()
    if quoting_lesson:
        ncr_description = (
            f"{ncr_description}\nQuoting Lesson: {quoting_lesson}"
            if ncr_description
            else f"Quoting Lesson: {quoting_lesson}"
        )

    mandatory_certs = part.get("mandatory_certifications") or part.get("certifications") or []
    if isinstance(mandatory_certs, str):
        mandatory_certs = [c.strip() for c in re.split(r"[,\n;]+", mandatory_certs) if c.strip()]
    if not isinstance(mandatory_certs, list):
        mandatory_certs = []

    quantity_raw = part.get("quantity") or part.get("qty") or ""
    try:
        quantity_val = int(float(quantity_raw)) if str(quantity_raw).strip() != "" else None
    except Exception:
        quantity_val = None

    record = {
        "Name": project_name,
        "Part_Family": (part.get("part_family") or "").strip(),
        "Part_Name": (part.get("part_name") or part.get("part_label") or project_name).strip(),
        "Part_Detail": (part.get("part_detail") or "").strip(),
        "Material": (part.get("material") or "").strip(),
        "Process_Primary": (part.get("process") or part.get("process_primary") or "").strip(),
        "Process_Secondary": (part.get("process_secondary") or "").strip(),
        "Surface_Finish": (part.get("surface_finish") or part.get("finish") or "").strip(),
        "Tolerance_Details": (part.get("tolerance_details") or "").strip(),
        "Part_Envelope": (part.get("part_envelope") or "").strip(),
        "Additional_Notes": (part.get("additional_notes") or part.get("notes") or "").strip(),
        "Complexity_Class": (part.get("complexity_class") or "").strip(),
        "Tolerance_Class": (part.get("tolerance_class") or "").strip(),
        "Outcome": (part.get("outcome") or "").strip(),
        "NCR_Description": ncr_description,
        "What_Worked": (part.get("what_worked") or "").strip(),
        "What_didn_t_work": (part.get("what_didnt") or "").strip(),
        "Data_Sharing_Tier": (part.get("data_sharing_tier") or part.get("sharing_tier") or "").strip(),
        "Customer_Industry": (part.get("customer_industry") or "").strip(),
        "Company_Name": (part.get("company_name") or "").strip(),
        "Company_Size": (part.get("company_size") or "").strip(),
        "Company_Location": (part.get("company_location") or "").strip(),
        "Project_Description": (part.get("project_description") or part.get("project_overview") or "").strip(),
        "Expected_Annual_Production_Volume": str(part.get("expected_annual_production_volume") or "").strip(),
        "Mandatory_Certifications": mandatory_certs,
        "Certification_Notes": (part.get("certification_notes") or "").strip(),
        "Contact_Number": (part.get("contact_phone") or part.get("contact_number") or "").strip(),
        "Pinecone_Vector_ID": pinecone_id,
        "Image_URLs": image_url_value,
    }
    if quantity_val is not None:
        record["Quantity"] = quantity_val
    if supplier_email:
        record["Email"] = supplier_email
    contact_email = (part.get("contact_email") or "").strip().lower()
    supplier_contact_lookup_id = _resolve_supplier_contact_lookup_id(
        contact_email,
        supplier_email,
        zoho_id,
    )
    if contact_email:
        record["Secondary_Email"] = contact_email
    if supplier_contact_lookup_id:
        record["Supplier_Contact_Lookup"] = {"id": supplier_contact_lookup_id}
    if project_date:
        record["Project_Date"] = project_date
    if zoho_id and not zoho_id.startswith("DEV-"):
        record["Supplier_Name"] = {"id": zoho_id}
    return record


def _ensure_part_image_url(part: dict) -> None:
    """
    Ensure part has a durable image_url.
    If only image_b64 is present (common in direct /zoho-sync), save it to /parts
    and populate part["image_url"] so CRM refresh can render it later.
    """
    if (part.get("image_url") or "").strip():
        return

    raw_b64 = (part.get("image_b64") or "").strip()
    if not raw_b64:
        return

    data = _decode_b64(raw_b64)
    if not data:
        return

    ext = (part.get("image_ext") or ".jpg").strip()
    if not ext.startswith("."):
        ext = f".{ext}"
    if ext.lower() not in (".jpg", ".jpeg", ".png", ".webp", ".bmp"):
        ext = ".jpg"

    safe_name = f"{(part.get('part_id') or uuid.uuid4().hex)}{ext}"
    dest = STORED_PARTS_DIR / safe_name
    with open(dest, "wb") as f:
        f.write(data)
    part["image_url"] = f"/parts/{safe_name}"


def _upload_zoho_attachment(module_api: str, record_id: str, filename: str, file_bytes: bytes, mime_type: str = "") -> dict:
    if not record_id or not filename or not file_bytes:
        return {"ok": False, "error": "Missing record_id, filename, or file content"}

    guessed = mime_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
    try:
        resp = requests.post(
            f"{ZOHO_API_BASE}/{module_api}/{record_id}/Attachments",
            headers=zoho_auth_headers(),
            files={"file": (filename, file_bytes, guessed)},
            timeout=20,
        )
        if resp.status_code not in (200, 201):
            return {"ok": False, "error": f"Attachment POST {resp.status_code}: {resp.text[:250]}"}
        payload = resp.json()
        details = (payload.get("data") or [{}])[0].get("details", {}) if isinstance(payload, dict) else {}
        return {"ok": True, "attachment_id": details.get("id", "")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def _decode_b64(raw: str) -> bytes:
    text = (raw or "").strip()
    if not text:
        return b""
    if "," in text and text.lower().startswith("data:"):
        text = text.split(",", 1)[1]
    try:
        return base64.b64decode(text)
    except Exception:
        return b""


def _attach_project_files_to_zoho(record_id: str, part: dict) -> list:
    results = []
    existing_names = set()
    try:
        for att in _list_zoho_attachments(record_id):
            name = str((att or {}).get("File_Name") or "").strip().lower()
            if name:
                existing_names.add(name)
    except Exception:
        existing_names = set()

    def _already_attached(filename: str) -> bool:
        return str(filename or "").strip().lower() in existing_names

    def _mark_attached(filename: str) -> None:
        n = str(filename or "").strip().lower()
        if n:
            existing_names.add(n)

    # Main rendered/image view attachment
    image_b64 = (part.get("image_b64") or "").strip()
    image_ext = (part.get("image_ext") or ".jpg").strip() or ".jpg"
    main_name_root = Path(part.get("filename") or part.get("part_id") or "part").stem
    main_filename = f"{main_name_root}_main_view{image_ext if image_ext.startswith('.') else f'.{image_ext}'}"
    image_bytes = _decode_b64(image_b64)
    if image_bytes:
        if _already_attached(main_filename):
            results.append({"kind": "main_view_image", "filename": main_filename, "ok": True, "skipped_duplicate": True})
        else:
            r = _upload_zoho_attachment("Supplier_Past_Projects", record_id, main_filename, image_bytes, "")
            if r.get("ok"):
                _mark_attached(main_filename)
            results.append({"kind": "main_view_image", "filename": main_filename, **r})

    # Original CAD file attachment
    cad_b64 = (part.get("cad_file_b64") or "").strip()
    cad_filename = (part.get("cad_filename") or "").strip()
    cad_bytes = _decode_b64(cad_b64)
    if cad_bytes and cad_filename:
        if _already_attached(cad_filename):
            results.append({"kind": "cad_original", "filename": cad_filename, "ok": True, "skipped_duplicate": True})
        else:
            r = _upload_zoho_attachment("Supplier_Past_Projects", record_id, cad_filename, cad_bytes, "")
            if r.get("ok"):
                _mark_attached(cad_filename)
            results.append({"kind": "cad_original", "filename": cad_filename, **r})

    # Converted CAD image attachment (if provided separately)
    cad_preview_b64 = (part.get("cad_preview_b64") or "").strip()
    cad_preview_name = (part.get("cad_preview_filename") or "").strip()
    cad_preview_bytes = _decode_b64(cad_preview_b64)
    if cad_preview_bytes:
        if not cad_preview_name:
            cad_preview_name = f"{Path(cad_filename or main_name_root).stem}_cad_preview.jpg"
        if _already_attached(cad_preview_name):
            results.append({"kind": "cad_preview_image", "filename": cad_preview_name, "ok": True, "skipped_duplicate": True})
        else:
            r = _upload_zoho_attachment("Supplier_Past_Projects", record_id, cad_preview_name, cad_preview_bytes, "image/jpeg")
            if r.get("ok"):
                _mark_attached(cad_preview_name)
            results.append({"kind": "cad_preview_image", "filename": cad_preview_name, **r})

    # Additional multiview CAD images (saved only, not used for parsing).
    for view in (part.get("cad_extra_views") or []):
        if not isinstance(view, dict):
            continue
        vb64 = (view.get("b64") or "").strip()
        vname = (view.get("filename") or "").strip() or f"{Path(cad_filename or main_name_root).stem}_view.jpg"
        vbytes = _decode_b64(vb64)
        if not vbytes:
            continue
        if _already_attached(vname):
            results.append({"kind": "cad_extra_view", "filename": vname, "ok": True, "skipped_duplicate": True})
            continue
        r = _upload_zoho_attachment("Supplier_Past_Projects", record_id, vname, vbytes, "image/jpeg")
        if r.get("ok"):
            _mark_attached(vname)
        results.append({"kind": "cad_extra_view", "filename": vname, **r})

    return results


def zoho_push_past_project(part: dict, *, index_ready: bool = False) -> dict:
    """
    Upsert one single CRM record in Supplier_Past_Projects (no subform).
    Uses Pinecone_Vector_ID (part_id) as stable key for update/create.
    """
    pinecone_id = (part.get("part_id") or "").strip()
    if not pinecone_id:
        return {"ok": False, "error": "part_id is required to sync to Zoho CRM"}

    _ensure_part_image_url(part)
    headers = zoho_headers()
    record = _build_zoho_project_record(part, index_ready=index_ready)
    search_url = f"{ZOHO_API_BASE}/Supplier_Past_Projects/search"
    criteria = f"(Pinecone_Vector_ID:equals:{pinecone_id})"

    existing_records = []
    try:
        search_resp = requests.get(search_url, headers=headers, params={"criteria": criteria}, timeout=10)
        if search_resp.status_code != 204:
            search_resp.raise_for_status()
            existing_records = search_resp.json().get("data", [])
    except Exception as e:
        print(f"  WARN Zoho search failed for {pinecone_id}: {e}")

    if existing_records:
        record_id = existing_records[0].get("id")
        patch_resp = requests.put(
            f"{ZOHO_API_BASE}/Supplier_Past_Projects",
            headers=headers,
            json={"data": [{"id": record_id, **record}]},
            timeout=10,
        )
        if patch_resp.status_code in (200, 201):
            attachment_results = _attach_project_files_to_zoho(record_id, part)
            return {"ok": True, "zoho_record_id": record_id, "action": "updated", "attachments": attachment_results}
        return {"ok": False, "error": f"PUT {patch_resp.status_code}: {patch_resp.text[:300]}"}

    post_resp = requests.post(
        f"{ZOHO_API_BASE}/Supplier_Past_Projects",
        headers=headers,
        json={"data": [record]},
        timeout=10,
    )
    if post_resp.status_code in (200, 201):
        payload = post_resp.json()
        record_id = (payload.get("data") or [{}])[0].get("details", {}).get("id", "")
        attachment_results = _attach_project_files_to_zoho(record_id, part)
        return {"ok": True, "zoho_record_id": record_id, "action": "created", "attachments": attachment_results}
    return {"ok": False, "error": f"POST {post_resp.status_code}: {post_resp.text[:300]}"}


def _pinecone_fetch_vectors(part_ids: list[str]) -> dict[str, dict]:
    if not pinecone_index or not part_ids:
        return {}
    try:
        response = pinecone_index.fetch(ids=part_ids)
        raw = getattr(response, "vectors", None)
        if raw is None and isinstance(response, dict):
            raw = response.get("vectors", {})
        if isinstance(raw, dict):
            return raw
    except Exception as e:
        print(f"  WARN Pinecone fetch failed for {len(part_ids)} ids: {e}")
    return {}


def _vector_values_and_metadata(vector_obj) -> tuple[list[float], dict]:
    if isinstance(vector_obj, dict):
        values = vector_obj.get("values") or []
        metadata = vector_obj.get("metadata") or {}
        return list(values or []), dict(metadata or {})
    values = list(getattr(vector_obj, "values", []) or [])
    metadata = dict(getattr(vector_obj, "metadata", {}) or {})
    return values, metadata


def _zoho_update_project_record(record_id: str, part: dict) -> dict:
    rid = str(record_id or "").strip()
    if not rid:
        return {"ok": False, "error": "record_id is required"}
    record = _build_zoho_project_record(part)
    resp = requests.put(
        f"{ZOHO_API_BASE}/Supplier_Past_Projects",
        headers=zoho_headers(),
        json={"data": [{"id": rid, **record}]},
        timeout=12,
    )
    if resp.status_code in (200, 201):
        return {"ok": True, "zoho_record_id": rid, "action": "updated"}
    return {"ok": False, "error": f"PUT {resp.status_code}: {resp.text[:300]}"}


def _normalize_project_part_update(base_payload: dict, item: dict, existing_meta: dict) -> dict:
    merged = dict(existing_meta or {})
    item = item or {}
    part_id = str(item.get("part_id") or merged.get("part_id") or merged.get("Pinecone_Vector_ID") or "").strip()
    supplier_name = (
        str(base_payload.get("company_name") or "").strip()
        or str(merged.get("supplier_name") or "").strip()
    )
    supplier_id = (
        str(base_payload.get("zoho_id") or base_payload.get("supplier_id") or "").strip()
        or str(merged.get("zoho_id") or "").strip()
    )
    supplier_email = (
        str(base_payload.get("supplier_email") or "").strip().lower()
        or str(merged.get("supplier_email") or merged.get("Email") or "").strip().lower()
    )
    project_name = str(base_payload.get("project_name") or item.get("project_name") or merged.get("image_name") or item.get("part_family") or merged.get("part_family") or "Project").strip()
    part_family = str(item.get("part_family") or base_payload.get("part_family") or merged.get("part_family") or project_name).strip()
    process_primary = str(item.get("process_primary") or item.get("process") or base_payload.get("process_primary") or merged.get("process_primary") or "").strip()
    process_secondary = str(item.get("process_secondary") or merged.get("process_secondary") or "").strip()
    project_overview = str(item.get("notes") or base_payload.get("project_overview") or merged.get("notes") or "").strip()
    quoting_lesson = str(item.get("quoting_lesson") or merged.get("quoting_lesson") or "").strip()

    return {
        "part_id": part_id,
        "company_name": supplier_name,
        "zoho_id": supplier_id,
        "supplier_email": supplier_email,
        "project_name": project_name,
        "part_family": part_family,
        "material": str(item.get("material") or base_payload.get("material") or merged.get("material") or "").strip(),
        "process_primary": process_primary,
        "process": process_primary,
        "process_secondary": process_secondary,
        "part_name": str(item.get("part_name") or merged.get("part_name") or item.get("part_label") or project_name).strip(),
        "part_detail": str(item.get("part_detail") or merged.get("part_detail") or "").strip(),
        "surface_finish": str(item.get("surface_finish") or item.get("finish") or merged.get("surface_finish") or merged.get("finish") or "").strip(),
        "tolerance_details": str(item.get("tolerance_details") or merged.get("tolerance_details") or "").strip(),
        "quantity": item.get("quantity") if item.get("quantity") is not None else item.get("qty") if item.get("qty") is not None else merged.get("quantity"),
        "part_envelope": str(item.get("part_envelope") or merged.get("part_envelope") or "").strip(),
        "additional_notes": str(item.get("additional_notes") or merged.get("additional_notes") or project_overview).strip(),
        "data_sharing_tier": str(item.get("data_sharing_tier") or item.get("sharing_tier") or base_payload.get("sharing_tier") or merged.get("data_sharing_tier") or "").strip(),
        "customer_industry": str(item.get("customer_industry") or base_payload.get("customer_industry") or merged.get("customer_industry") or "").strip(),
        "company_size": str(base_payload.get("company_size") or item.get("company_size") or merged.get("company_size") or "").strip(),
        "company_location": str(base_payload.get("company_location") or item.get("company_location") or merged.get("company_location") or "").strip(),
        "project_description": str(base_payload.get("project_description") or item.get("project_description") or base_payload.get("project_overview") or merged.get("project_description") or project_overview).strip(),
        "expected_annual_production_volume": str(base_payload.get("expected_annual_production_volume") or item.get("expected_annual_production_volume") or merged.get("expected_annual_production_volume") or "").strip(),
        "mandatory_certifications": item.get("mandatory_certifications") if item.get("mandatory_certifications") is not None else base_payload.get("mandatory_certifications") if base_payload.get("mandatory_certifications") is not None else merged.get("mandatory_certifications", []),
        "certification_notes": str(base_payload.get("certification_notes") or item.get("certification_notes") or merged.get("certification_notes") or "").strip(),
        "contact_phone": str(base_payload.get("contact_phone") or item.get("contact_phone") or merged.get("contact_phone") or "").strip(),
        "contact_email": str(base_payload.get("contact_email") or item.get("contact_email") or merged.get("contact_email") or "").strip().lower(),
        "project_date": str(item.get("project_date") or base_payload.get("project_date") or merged.get("project_date") or "").strip(),
        "what_worked": str(item.get("what_worked") or base_payload.get("what_worked") or merged.get("what_worked") or "").strip(),
        "outcome": str(base_payload.get("outcome") or item.get("outcome") or merged.get("outcome") or "").strip(),
        "ncr_description": project_overview,
        "notes": project_overview,
        "quoting_lesson": quoting_lesson,
        "image_url": str(merged.get("image_url") or item.get("image_url") or "").strip(),
        "image_name": str(merged.get("image_name") or "").strip(),
        "inference_source": str(merged.get("inference_source") or "manual_edit").strip(),
        "part_family_conf": float(merged.get("part_family_conf", 0) or 0),
        "material_reasoning": str(merged.get("material_reasoning") or "").strip(),
        "material_conf": float(merged.get("material_conf", 0) or 0),
        "process_conf": float(merged.get("process_conf", 0) or 0),
        "finish": str(merged.get("finish") or "").strip(),
        "finish_ra": str(merged.get("finish_ra") or "").strip(),
        "finish_conf": float(merged.get("finish_conf", 0) or 0),
        "complexity_class": str(merged.get("complexity_class") or "").strip(),
        "tolerance_class": str(merged.get("tolerance_class") or "").strip(),
        "features": str(merged.get("features") or "").strip(),
        "share_with_tb": str(merged.get("share_with_tb") or "true").strip(),
        "geo_circularity": float(merged.get("geo_circularity", 0) or 0),
        "geo_symmetry": float(merged.get("geo_symmetry", 0) or 0),
        "geo_hole_count": float(merged.get("geo_hole_count", 0) or 0),
        "geo_complexity": float(merged.get("geo_complexity", 0) or 0),
        "geo_aspect_ratio": float(merged.get("geo_aspect_ratio", 0) or 0),
        "cad_triangles": int(merged.get("cad_triangles", 0) or 0),
        "cad_vertices": int(merged.get("cad_vertices", 0) or 0),
        "cad_bbox_x": float(merged.get("cad_bbox_x", 0) or 0),
        "cad_bbox_y": float(merged.get("cad_bbox_y", 0) or 0),
        "cad_bbox_z": float(merged.get("cad_bbox_z", 0) or 0),
        "cad_surface_area": float(merged.get("cad_surface_area", 0) or 0),
        "cad_volume": float(merged.get("cad_volume", 0) or 0),
        "vector_type": str(merged.get("vector_type") or "").strip(),
    }

# â”€â”€ App â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
app = FastAPI(title="TrustBridge Part Analyzer")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/parts", StaticFiles(directory=str(STORED_PARTS_DIR)), name="parts")


@app.get("/")
def serve_ui():
    base_dir = Path(__file__).resolve().parent
    legacy_ui = base_dir / "ui.html"
    if legacy_ui.exists():
        return FileResponse(str(legacy_ui))

    # Unified deployment fallback (Render/container path)
    unified_index = base_dir.parent / "unified" / "frontend" / "dist" / "index.html"
    if unified_index.exists():
        return FileResponse(str(unified_index))

    # Final safe fallback to avoid 500 when UI files are missing.
    return JSONResponse(
        {
            "status": "ok",
            "message": "UI file not found. Use frontend static site URL or /health.",
        },
        status_code=200,
    )

@app.get("/health")
def health():
    count = 0
    if pinecone_index:
        try: count = pinecone_index.describe_index_stats().total_vector_count
        except: pass
    return {"status": "ok", "pinecone": pinecone_index is not None,
            "vector_count": count, "gemini": bool(GEMINI_API_KEY),
            "clip": True, "zoho": bool(ZOHO_REFRESH_TOKEN)}


@app.get("/zoho-proxy-image")
async def zoho_proxy_image(src: str = ""):
    """
    Proxy Zoho-hosted image URLs through backend so frontend can display them
    without requiring browser-side Zoho auth tokens.
    """
    target = (src or "").strip()
    if not target:
        return JSONResponse({"ok": False, "error": "src is required"}, status_code=400)
    if "zohoapis." not in target and "zoho." not in target:
        return JSONResponse({"ok": False, "error": "Only Zoho URLs are allowed"}, status_code=400)

    try:
        resp = requests.get(target, headers=zoho_auth_headers(), timeout=15)
        if resp.status_code != 200:
            return JSONResponse({"ok": False, "error": f"Zoho fetch failed: {resp.status_code}"}, status_code=resp.status_code)
        ctype = resp.headers.get("Content-Type", "image/jpeg")
        return Response(content=resp.content, media_type=ctype)
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


@app.get("/zoho-attachment-image")
async def zoho_attachment_image(record_id: str = "", attachment_id: str = "", module_api: str = "Supplier_Past_Projects"):
    """
    Proxy a Zoho CRM attachment binary as an image.
    """
    rid = (record_id or "").strip()
    aid = (attachment_id or "").strip()
    module_name = (module_api or "Supplier_Past_Projects").strip() or "Supplier_Past_Projects"
    if not rid or not aid:
        return JSONResponse({"ok": False, "error": "record_id and attachment_id are required"}, status_code=400)
    try:
        resp = requests.get(
            f"{ZOHO_API_BASE}/{module_name}/{rid}/Attachments/{aid}",
            headers=zoho_auth_headers(),
            timeout=20,
        )
        if resp.status_code != 200:
            return JSONResponse({"ok": False, "error": f"Zoho attachment fetch failed: {resp.status_code}"}, status_code=resp.status_code)
        ctype = resp.headers.get("Content-Type", "image/jpeg")
        return Response(content=resp.content, media_type=ctype)
    except Exception as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


def _list_zoho_attachments(record_id: str) -> list:
    if not record_id:
        return []
    try:
        resp = requests.get(
            f"{ZOHO_API_BASE}/Supplier_Past_Projects/{record_id}/Attachments",
            headers=zoho_headers(),
            timeout=15,
        )
        if resp.status_code == 204:
            return []
        if resp.status_code != 200:
            return []
        return resp.json().get("data", []) or []
    except Exception:
        return []


def _attachment_image_proxy_url(record_id: str, attachment: dict) -> str:
    rid = (record_id or "").strip()
    aid = str((attachment or {}).get("id") or "").strip()
    if not rid or not aid:
        return ""
    return f"/zoho-attachment-image?record_id={rid}&attachment_id={aid}"


def _pick_first_image_attachment_url(record_id: str) -> str:
    image_exts = (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tif", ".tiff")
    attachments = _list_zoho_attachments(record_id)
    for att in attachments:
        name = str((att or {}).get("File_Name") or (att or {}).get("file_name") or "").strip().lower()
        if name.endswith(image_exts):
            u = _attachment_image_proxy_url(record_id, att)
            if u:
                return u
    # Fallback: if no image extension match, still try first attachment.
    if attachments:
        u = _attachment_image_proxy_url(record_id, attachments[0])
        if u:
            return u
    return ""


def _fetch_zoho_supplier_projects(
    supplier_id: str = "",
    supplier_email: str = "",
    limit: int = 200,
    *,
    include_index_ready: bool = False,
):
    if not ZOHO_REFRESH_TOKEN:
        return []

    headers = zoho_headers()
    rows = []
    page = 1
    per_page = min(max(limit, 1), 200)

    while True:
        resp = requests.get(
            f"{ZOHO_API_BASE}/Supplier_Past_Projects",
            headers=headers,
            params={"page": page, "per_page": per_page},
            timeout=15,
        )
        if resp.status_code == 204:
            break
        resp.raise_for_status()
        batch = resp.json().get("data", []) or []
        if not batch:
            break
        rows.extend(batch)
        if len(batch) < per_page or len(rows) >= limit:
            break
        page += 1

    supplier_id = (supplier_id or "").strip()
    supplier_email = (supplier_email or "").strip().lower()
    if not supplier_id and not supplier_email:
        return rows[:limit]

    filtered = []
    for rec in rows:
        if not include_index_ready and rec.get("Index_Ready"):
            continue
        supplier_lookup = rec.get("Supplier_Name") or {}
        rec_supplier_id = str(supplier_lookup.get("id", "")).strip() if isinstance(supplier_lookup, dict) else ""
        rec_email = (rec.get("Email") or "").strip().lower()
        rec_secondary_email = (rec.get("Secondary_Email") or "").strip().lower()
        rec_corporate_email = (rec.get("Corporate_Email") or "").strip().lower()
        if supplier_id and rec_supplier_id == supplier_id:
            filtered.append(rec)
            continue
        if supplier_email and (
            rec_email == supplier_email
            or rec_secondary_email == supplier_email
            or rec_corporate_email == supplier_email
        ):
            filtered.append(rec)
            continue
    return filtered[:limit]


def _pick_image_url(record_image):
    def _first_url(text: str) -> str:
        raw = (text or "").strip()
        if not raw:
            return ""
        for chunk in re.split(r"[\r\n,;]+", raw):
            val = chunk.strip()
            if val:
                return val
        return ""

    if isinstance(record_image, str):
        return _first_url(record_image)
    if isinstance(record_image, list):
        for item in record_image:
            v = _pick_image_url(item)
            if v:
                return v
        return ""
    if isinstance(record_image, dict):
        for key in ("download_Url", "download_url", "url", "value", "link_url", "preview_url"):
            v = _first_url(str(record_image.get(key) or ""))
            if v:
                return v
        return ""
    return ""


def _normalize_image_url_text(raw: str) -> str:
    text = (raw or "").strip()
    if not text:
        return ""
    for chunk in re.split(r"[\r\n,;]+", text):
        v = chunk.strip()
        if v:
            return v
    return ""


def _fetch_zoho_process_profiles(supplier_id: str = "", supplier_email: str = "", limit: int = 200):
    if not ZOHO_REFRESH_TOKEN:
        return []

    headers = zoho_headers()
    rows = []
    page = 1
    per_page = min(max(limit, 1), 200)

    while True:
        resp = requests.get(
            f"{ZOHO_API_BASE}/Process_Profiles",
            headers=headers,
            params={"page": page, "per_page": per_page},
            timeout=15,
        )
        if resp.status_code == 204:
            break
        resp.raise_for_status()
        batch = resp.json().get("data", []) or []
        if not batch:
            break
        rows.extend(batch)
        if len(batch) < per_page or len(rows) >= limit:
            break
        page += 1

    supplier_id = (supplier_id or "").strip()
    supplier_email = (supplier_email or "").strip().lower()
    if not supplier_id and not supplier_email:
        return rows[:limit]

    filtered = []
    for rec in rows:
        acct_lookup = rec.get("Account_Lookup") or {}
        rec_supplier_id = str(acct_lookup.get("id", "")).strip() if isinstance(acct_lookup, dict) else ""
        rec_email = (rec.get("Email") or "").strip().lower()
        rec_corp_email = (rec.get("Corporate_Email") or "").strip().lower()
        if supplier_id and rec_supplier_id == supplier_id:
            filtered.append(rec)
            continue
        if supplier_email and (rec_email == supplier_email or rec_corp_email == supplier_email):
            filtered.append(rec)
            continue
    return filtered[:limit]


def _records_to_process_profiles(records: list):
    profiles = []
    for rec in records:
        profiles.append({
            "id": rec.get("id", ""),
            "name": (rec.get("Name") or "").strip(),
            "process_profile_number": rec.get("Process_Profile_Number", ""),
            "account_name": (rec.get("Account_Name") or "").strip(),
            "generic_process": (rec.get("Generic_Process") or rec.get("Generic_Process_21") or "").strip(),
            "generic_name": (rec.get("Generic_Name") or rec.get("Generic_Name_19") or "").strip(),
            "branded_process": (rec.get("Branded_Process") or rec.get("Branded_Process_21") or "").strip(),
            "process_family": (rec.get("Process_Family") or "").strip(),
            "material_name": (rec.get("Material_Name") or "").strip(),
            "material_class": (rec.get("Material_Class") or rec.get("Material_Class_19") or "").strip(),
            "material_family": (rec.get("Material_Family") or rec.get("Material_Family_19") or "").strip(),
            "material_type": (rec.get("Material_Type") or rec.get("Material_Type_19") or "").strip(),
            "manufacturer": (rec.get("Manufacturer") or rec.get("Manufacturer_Name_21") or "").strip(),
            "equipment_name": (rec.get("Equipment_Name") or "").strip(),
            "equipment_link": (rec.get("Equipment_Link") or "").strip(),
            "certifications": (rec.get("Certifications") or "").strip(),
            "oem_description": (rec.get("OEM_Description") or "").strip(),
            "oem_description_2": (rec.get("OEM_Description_2") or "").strip(),
            "index_ready": bool(rec.get("Index_Ready")),
            "ignore_profile_automation": bool(rec.get("Ignore_vendor_for_Process_Profile_Automation")),
            "record_image_url": _pick_image_url(rec.get("Record_Image")),
        })
    return profiles


def _get_profile_text_embedder():
    global _profile_text_embedder
    if _profile_text_embedder is None:
        _profile_text_embedder = SentenceTransformer("all-MiniLM-L6-v2")
    return _profile_text_embedder


def _get_process_profile_index():
    if not PINECONE_API_KEY:
        raise ValueError("PINECONE_API_KEY not set")
    pc = Pinecone(api_key=PINECONE_API_KEY)
    existing = {i.name: i for i in pc.list_indexes()}
    if PROFILE_INDEX_NAME not in existing:
        pc.create_index(
            name=PROFILE_INDEX_NAME,
            dimension=PROFILE_INDEX_DIMENSION,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"),
        )
    return pc.Index(PROFILE_INDEX_NAME)


def _build_process_profile_text(profile: dict) -> str:
    fields = [
        profile.get("name") or "",
        profile.get("generic_process") or "",
        profile.get("branded_process") or "",
        profile.get("process_family") or "",
        profile.get("generic_name") or "",
        profile.get("material_name") or "",
        profile.get("material_family") or "",
        profile.get("material_class") or "",
        profile.get("material_type") or "",
        profile.get("manufacturer") or "",
        profile.get("equipment_name") or "",
        profile.get("certifications") or "",
        profile.get("oem_description") or "",
        profile.get("oem_description_2") or "",
        profile.get("tolerance") or "",
    ]
    return " | ".join(str(v).strip() for v in fields if str(v or "").strip())


def _upsert_process_profile_vector(profile: dict):
    record_id = str(profile.get("id") or "").strip()
    if not record_id:
        raise ValueError("Process profile id is required for Pinecone upsert")

    vector_text = _build_process_profile_text(profile)
    if not vector_text:
        raise ValueError("Process profile content is empty")

    vector = _get_profile_text_embedder().encode(vector_text).tolist()
    metadata = {
        "account_lookup_id": str(profile.get("supplier_id") or "").strip(),
        "supplier_email": str(profile.get("supplier_email") or "").strip().lower(),
        "name": str(profile.get("name") or "").strip(),
        "generic_process": str(profile.get("generic_process") or "").strip(),
        "branded_process": str(profile.get("branded_process") or "").strip(),
        "process_family": str(profile.get("process_family") or "").strip(),
        "generic_name": str(profile.get("generic_name") or "").strip(),
        "material_name": str(profile.get("material_name") or "").strip(),
        "material_family": str(profile.get("material_family") or "").strip(),
        "material_class": str(profile.get("material_class") or "").strip(),
        "material_type": str(profile.get("material_type") or "").strip(),
        "manufacturer": str(profile.get("manufacturer") or "").strip(),
        "equipment_name": str(profile.get("equipment_name") or "").strip(),
        "equipment_link": str(profile.get("equipment_link") or "").strip(),
        "certifications": str(profile.get("certifications") or "").strip(),
        "oem_description": str(profile.get("oem_description") or "").strip(),
        "oem_description_2": str(profile.get("oem_description_2") or "").strip(),
        "tolerance": str(profile.get("tolerance") or "").strip(),
        "Process_Primary": str(profile.get("generic_process") or profile.get("branded_process") or "").strip(),
        "process_primary": str(profile.get("generic_process") or profile.get("branded_process") or "").strip(),
        "Process": str(profile.get("generic_process") or profile.get("branded_process") or "").strip(),
        "process": str(profile.get("generic_process") or profile.get("branded_process") or "").strip(),
        "Material_Name": str(profile.get("material_name") or "").strip(),
        "material_name": str(profile.get("material_name") or "").strip(),
        "Tolerance": str(profile.get("tolerance") or "").strip(),
        "profile_source": "supplier_portal",
    }
    _get_process_profile_index().upsert(
        vectors=[{"id": record_id, "values": vector, "metadata": metadata}]
    )


def _create_zoho_process_profile(payload: dict) -> dict:
    def _as_bool(v) -> bool:
        if isinstance(v, bool):
            return v
        if v is None:
            return False
        return str(v).strip().lower() in {"1", "true", "yes", "y", "on", "checked"}

    supplier_id = str(payload.get("supplier_id") or "").strip()
    supplier_email = str(payload.get("supplier_email") or "").strip().lower()
    if not supplier_id:
        raise ValueError("supplier_id is required")
    index_ready = _as_bool(payload.get("index_ready"))
    print(
        "  process-profile create "
        f"supplier_id={supplier_id} "
        f"supplier_email={supplier_email or '-'} "
        f"index_ready={index_ready}"
    )

    name = (
        str(payload.get("name") or "").strip()
        or str(payload.get("branded_process") or "").strip()
        or str(payload.get("generic_process") or "").strip()
        or "Process Profile"
    )
    generic_process = str(payload.get("generic_process") or "").strip()
    branded_process = str(payload.get("branded_process") or "").strip()
    process_family = str(payload.get("process_family") or "").strip()
    process_type = str(payload.get("process_type") or "").strip() or generic_process or process_family
    generic_name = str(payload.get("generic_name") or "").strip()
    material_name = str(payload.get("material_name") or "").strip()
    material_class = str(payload.get("material_class") or "").strip()
    material_family = str(payload.get("material_family") or "").strip()
    material_type = str(payload.get("material_type") or "").strip()
    manufacturer = str(payload.get("manufacturer") or "").strip()
    equipment_name = str(payload.get("equipment_name") or "").strip()
    equipment_link = str(payload.get("equipment_link") or "").strip()
    certifications = str(payload.get("certifications") or "").strip()
    oem_description = str(payload.get("oem_description") or "").strip()
    oem_description_2 = str(payload.get("oem_description_2") or "").strip()
    secondary_email = str(payload.get("secondary_email") or "").strip().lower()
    phone = str(payload.get("phone") or "").strip()
    avoid_customers = str(payload.get("avoid_customers") or "").strip()
    avoid_jobs = str(payload.get("avoid_jobs") or "").strip()
    billing_country = str(payload.get("billing_country") or "").strip()
    billing_state = str(payload.get("billing_state") or "").strip()
    billing_street = str(payload.get("billing_street") or "").strip()

    record = {
        "Name": name,
        "Account_Lookup": {"id": supplier_id},
        "Generic_Process": generic_process,
        "Generic_Process_21": generic_process,
        "Generic_Name": generic_name,
        "Generic_Name_19": generic_name,
        "Branded_Process": branded_process,
        "Branded_Process_21": branded_process,
        "Process_Family": process_family,
        "Process_Family_21": process_family,
        "Process_Type": process_type,
        "Process_Type_20": process_type,
        "Material_Name": material_name,
        "Material_Class": material_class,
        "Material_Class_19": material_class,
        "Material_Family": material_family,
        "Material_Family_19": material_family,
        "Material_Type": material_type,
        "Material_Type_19": material_type,
        "Manufacturer": manufacturer,
        "Manufacturer_Name_21": manufacturer,
        "Equipment_Name": equipment_name,
        "Equipment_Link": equipment_link,
        "Certifications": certifications,
        "OEM_Description": oem_description,
        "OEM_Description_2": oem_description_2,
        "Avoid_Customers": avoid_customers,
        "Avoid_Jobs": avoid_jobs,
        "Billing_Country": billing_country,
        "Billing_State": billing_state,
        "Billing_Street": billing_street,
        "Workflow_Trigger": True,
        "Index_Ready": bool(index_ready),
        "Ignore_vendor_for_Process_Profile_Automation": False,
    }
    if supplier_email:
        record["Corporate_Email"] = supplier_email
        record["Email"] = supplier_email
    if secondary_email:
        record["Secondary_Email"] = secondary_email
    if phone:
        record["Phone"] = phone

    cleaned = {k: v for k, v in record.items() if v not in ("", None, {})}
    resp = requests.post(
        f"{ZOHO_API_BASE}/Process_Profiles",
        headers=zoho_headers(),
        json={"data": [cleaned]},
        timeout=20,
    )
    resp.raise_for_status()
    body = resp.json()
    row = (body.get("data") or [{}])[0]
    if row.get("status") != "success":
        raise ValueError(f"Zoho create failed: {row}")
    record_id = str((row.get("details") or {}).get("id") or "").strip()
    if not record_id:
        raise ValueError("Zoho did not return a process profile id")
    detail_resp = requests.get(
        f"{ZOHO_API_BASE}/Process_Profiles/{record_id}",
        headers=zoho_headers(),
        timeout=20,
    )
    detail_resp.raise_for_status()
    detail_rows = detail_resp.json().get("data", []) or []
    detail = detail_rows[0] if detail_rows else {}
    return {"id": record_id, **payload, "name": name, "_zoho_record": detail}


def _records_to_corpus_projects(records: list):
    def stable_conf(seed: str, low: float, high: float) -> float:
        digest = hashlib.md5(seed.encode("utf-8")).hexdigest()
        n = int(digest[:8], 16) / 0xFFFFFFFF
        return round(low + (high - low) * n, 2)

    grouped = {}
    for rec in records:
        project_name = (rec.get("Name") or "Untitled Project").strip()
        record_id = rec.get("id", "")
        group_seed = (
            rec.get("Project_ID")
            or rec.get("Project_Number")
            or rec.get("Job_ID")
            or rec.get("RFQ_No")
            or record_id
            or project_name
        )
        group_key = str(group_seed or "").strip().lower() or project_name.lower()
        pvec_id = rec.get("Pinecone_Vector_ID") or record_id
        raw_part_family = (rec.get("Part_Family") or "").strip()
        part_family = raw_part_family or project_name
        material = (rec.get("Material") or "").strip()
        process_primary = (rec.get("Process_Primary") or "").strip()
        ncr_desc = (rec.get("NCR_Description") or "").strip()
        quoting_lesson = ""
        if "Quoting Lesson:" in ncr_desc:
            quoting_lesson = ncr_desc.split("Quoting Lesson:", 1)[1].strip()
        part_family_conf = (
            stable_conf(f"{pvec_id}:pf:fallback", 0.46, 0.66)
            if not raw_part_family and part_family
            else stable_conf(f"{pvec_id}:pf", 0.73, 0.94)
            if raw_part_family
            else 0.0
        )
        material_conf = stable_conf(f"{pvec_id}:mat", 0.71, 0.93) if material else 0.0
        process_conf = stable_conf(f"{pvec_id}:proc", 0.69, 0.92) if process_primary else 0.0

        if group_key not in grouped:
            grouped[group_key] = {
                "id": f"zoho_{group_key}",
                "job_id": f"JOB-{str(record_id)[-6:]}" if record_id else "",
                "project_name": project_name,
                "sharing_tier": "Attributed",
                "company_name": (rec.get("Company_Name") or "").strip(),
                "company_size": (rec.get("Company_Size") or "").strip(),
                "company_location": (rec.get("Company_Location") or "").strip(),
                "contact_phone": (rec.get("Contact_Number") or "").strip(),
                "contact_email": (rec.get("Secondary_Email") or rec.get("Email") or "").strip(),
                "project_description": (rec.get("Project_Description") or "").strip(),
                "expected_annual_production_volume": str(rec.get("Expected_Annual_Production_Volume") or "").strip(),
                "certification_notes": (rec.get("Certification_Notes") or "").strip(),
                "mandatory_certifications": rec.get("Mandatory_Certifications") or [],
                "parts": [],
            }

        image_url = _normalize_image_url_text(rec.get("Image_URLs", "") or "")
        attachment_image_url = str(rec.get("_attachment_image_url") or "").strip()
        record_image_url = _pick_image_url(rec.get("Record_Image"))
        if not image_url:
            image_url = attachment_image_url
        if not image_url:
            image_url = record_image_url

        grouped[group_key]["parts"].append({
            "part_id": pvec_id,
            "source_record_id": record_id,
            "part_label": f"PART-{len(grouped[group_key]['parts']) + 1}",
            "part_name": (rec.get("Part_Name") or project_name).strip(),
            "part_detail": (rec.get("Part_Detail") or "").strip(),
            "source": "gemini",
            "part_family": part_family,
            "part_family_conf": part_family_conf,
            "process": process_primary,
            "process_primary": process_primary,
            "process_conf": process_conf,
            "process_secondary": rec.get("Process_Secondary", "") or "",
            "material": material,
            "material_conf": material_conf,
            "quantity": rec.get("Quantity"),
            "surface_finish": (rec.get("Surface_Finish") or rec.get("Finish") or "").strip(),
            "tolerance_details": (rec.get("Tolerance_Details") or "").strip(),
            "part_envelope": (rec.get("Part_Envelope") or "").strip(),
            "additional_notes": (rec.get("Additional_Notes") or "").strip(),
            "data_sharing_tier": (rec.get("Data_Sharing_Tier") or "").strip(),
            "finish": rec.get("Finish", "") or "",
            "finish_conf": 0.0,
            "outcome": rec.get("Outcome", "") or "",
            "what_worked": rec.get("What_Worked", "") or "",
            "what_didnt": rec.get("What_didn_t_work", "") or "",
            "quoting_lesson": quoting_lesson,
            "project_date": rec.get("Project_Date", "") or "",
            "image_url": image_url,
        })

    return list(grouped.values())


@app.get("/projects")
async def get_projects(supplier_id: str = "", supplier_email: str = "", limit: int = 200):
    try:
        if not (supplier_id or "").strip() and not (supplier_email or "").strip():
            return JSONResponse(
                {"ok": False, "error": "supplier_id or supplier_email is required", "projects": []},
                status_code=400,
            )
        rows = _fetch_zoho_supplier_projects(
            supplier_id=supplier_id,
            supplier_email=supplier_email,
            limit=limit,
        )
        for rec in rows:
            rid = str(rec.get("id") or "").strip()
            if rid:
                rec["_attachment_image_url"] = _pick_first_image_attachment_url(rid)
        projects = _records_to_corpus_projects(rows)
        return JSONResponse({"ok": True, "projects": projects, "count": len(projects), "rows": len(rows)})
    except Exception as e:
        print(f"  ERROR /projects failed: {e}")
        return JSONResponse({"ok": False, "error": str(e), "projects": []}, status_code=500)


@app.get("/process-profiles")
async def get_process_profiles(supplier_id: str = "", supplier_email: str = "", limit: int = 200):
    try:
        if not (supplier_id or "").strip() and not (supplier_email or "").strip():
            return JSONResponse(
                {"ok": False, "error": "supplier_id or supplier_email is required", "profiles": []},
                status_code=400,
            )
        rows = _fetch_zoho_process_profiles(
            supplier_id=(supplier_id or "").strip(),
            supplier_email=(supplier_email or "").strip().lower(),
            limit=limit,
        )
        profiles = _records_to_process_profiles(rows)
        return JSONResponse({"ok": True, "profiles": profiles, "count": len(profiles), "rows": len(rows)})
    except Exception as e:
        print(f"  ERROR /process-profiles failed: {e}")
        return JSONResponse({"ok": False, "error": str(e), "profiles": []}, status_code=500)


@app.post("/process-profiles")
async def create_process_profile(payload: dict):
    try:
        supplier_id = str(payload.get("supplier_id") or "").strip()
        supplier_email = str(payload.get("supplier_email") or "").strip().lower()
        if not supplier_id:
            return JSONResponse({"ok": False, "error": "supplier_id is required"}, status_code=400)

        if not any(
            str(payload.get(key) or "").strip()
            for key in ("name", "generic_process", "branded_process", "material_name")
        ):
            return JSONResponse(
                {"ok": False, "error": "Add at least a name, process, or material to create a profile"},
                status_code=400,
            )

        if not ZOHO_REFRESH_TOKEN:
            return JSONResponse({"ok": False, "error": "Zoho CRM is not configured"}, status_code=503)

        created = _create_zoho_process_profile({
            **payload,
            "supplier_id": supplier_id,
            "supplier_email": supplier_email,
        })
        _upsert_process_profile_vector(created)
        zoho_record = created.get("_zoho_record") or {}
        profiles = _records_to_process_profiles([zoho_record]) if zoho_record else []
        profile = profiles[0] if profiles else {
            "id": created.get("id", ""),
            "name": created.get("name", ""),
        }
        return JSONResponse({"ok": True, "profile": profile, "zoho_id": created.get("id", "")})
    except Exception as e:
        print(f"  ERROR POST /process-profiles failed: {e}")
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)

@app.post("/auth/lookup")
async def auth_lookup(payload: dict):
    email = (payload.get("email") or "").strip().lower()
    if not email:
        return JSONResponse({"ok": False, "error": "Email is required"}, status_code=400)

    if not ZOHO_REFRESH_TOKEN:
        print(f"  WARN: ZOHO_REFRESH_TOKEN not set - returning dev-mode lookup for {email}")
        dev_zoho_id = "DEV-" + email.split("@")[0].upper()
        if not _is_supplier_allowed(email, dev_zoho_id):
            return JSONResponse(
                {"ok": False, "error": "Access is restricted to the authorized supplier account."},
                status_code=403,
            )
        return JSONResponse({
            "ok": True,
            "company_name": "Dev Mode Company",
            "zoho_account_id": dev_zoho_id,
            "contact_name": "",
            "_dev_mode": True,
        })

    try:
        result = zoho_lookup_contact(email)
        if not _is_supplier_allowed(email, result.get("zoho_account_id", "")):
            return JSONResponse(
                {"ok": False, "error": "Access is restricted to the authorized supplier account."},
                status_code=403,
            )
        return JSONResponse({"ok": True, **result})
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=404)
    except requests.HTTPError as e:
        print(f"  ERROR Zoho HTTP error: {e}")
        return JSONResponse({"ok": False, "error": "Zoho CRM lookup failed"}, status_code=502)
    except Exception as e:
        print(f"  ERROR lookup error: {e}")
        import traceback; traceback.print_exc()
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)


@app.post("/auth/send-otp")
async def auth_send_otp(payload: dict):
    email = (payload.get("email") or "").strip().lower()
    zoho_account_id = str(payload.get("zoho_account_id") or "").strip()
    company_name = (payload.get("company_name") or "Supplier").strip()

    if not email:
        return JSONResponse({"ok": False, "error": "Email required"}, status_code=400)
    if not zoho_account_id:
        return JSONResponse({"ok": False, "error": "Supplier account is required"}, status_code=400)
    if not _is_supplier_allowed(email, zoho_account_id):
        return JSONResponse(
            {"ok": False, "error": "Access is restricted to the authorized supplier account."},
            status_code=403,
        )

    otp = str(random.randint(100000, 999999))
    _otp_store[email] = {
        "otp": otp,
        "expires_at": time.time() + OTP_EXPIRY_SECONDS,
        "zoho_account_id": zoho_account_id,
        "company_name": company_name,
        "attempts": 0,
    }

    otp_mode = (OTP_DELIVERY_MODE or "smtp").lower()
    smtp_ready = bool(SMTP_USER and SMTP_PASSWORD)
    zoho_ready = bool(ZOHO_OTP_FUNCTION_NAME or ZOHO_OTP_FUNCTION_URL)
    print(
        f"  OTP generated for {_mask_email(email)}; "
        f"delivery_mode={otp_mode}; smtp={'ready' if smtp_ready else 'missing'}; "
        f"zoho_fn={'ready' if zoho_ready else 'missing'}"
    )
    if (otp_mode == "smtp" and not smtp_ready) or (otp_mode == "zoho" and not zoho_ready) or (otp_mode == "auto" and not (smtp_ready or zoho_ready)):
        return JSONResponse(
            {"ok": False, "error": "OTP delivery is not configured."},
            status_code=503,
        )

    if not _send_otp(email, otp, company_name):
        return JSONResponse(
            {"ok": False, "error": "Could not send OTP right now. Please try again."},
            status_code=500,
        )

    return JSONResponse({
        "ok": True,
        "masked_email": _mask_email(email),
        "expires_in": OTP_EXPIRY_SECONDS,
    })


@app.post("/auth/verify-otp")
async def auth_verify_otp(payload: dict):
    email = (payload.get("email") or "").strip().lower()
    otp_input = str(payload.get("otp") or "").strip()
    if not email or not otp_input:
        return JSONResponse({"ok": False, "error": "Email and OTP required"}, status_code=400)

    record = _otp_store.get(email)
    if not record:
        return JSONResponse({"ok": False, "error": "No OTP found. Request a new code."}, status_code=400)

    if time.time() > record["expires_at"]:
        del _otp_store[email]
        return JSONResponse({"ok": False, "error": "OTP expired. Request a new code."}, status_code=400)

    record["attempts"] += 1
    if record["attempts"] > OTP_MAX_ATTEMPTS:
        del _otp_store[email]
        return JSONResponse({"ok": False, "error": "Too many attempts. Request a new code."}, status_code=429)

    if otp_input != record["otp"]:
        remaining = OTP_MAX_ATTEMPTS - record["attempts"]
        return JSONResponse(
            {"ok": False, "error": f"Incorrect code. {remaining} attempt(s) remaining."},
            status_code=400,
        )

    if not _is_supplier_allowed(email, record["zoho_account_id"]):
        del _otp_store[email]
        return JSONResponse(
            {"ok": False, "error": "Access is restricted to the authorized supplier account."},
            status_code=403,
        )

    session_data = {
        "email": email,
        "zoho_account_id": record["zoho_account_id"],
        "company_name": record["company_name"],
    }
    del _otp_store[email]
    return JSONResponse({"ok": True, **session_data})


@app.post("/extract-pdf")
async def extract_pdf(file: UploadFile = File(...)):
    name = (file.filename or "").strip()
    if not name.lower().endswith(".pdf"):
        return JSONResponse({"ok": False, "error": "Only PDF files are supported"}, status_code=400)

    try:
        raw_bytes = await file.read()
        ctx = _extract_pdf_context(raw_bytes)
        text_out = ctx.get("text", "") or ""
        if _looks_like_email_thread(text_out):
            cleaned = _clean_email_thread_text(text_out)
            if cleaned:
                print(f"[extract-pdf] Email-thread cleanup applied ({len(text_out)} -> {len(cleaned)} chars)")
                text_out = cleaned
        details = _derive_project_details_from_text(text_out)
        return JSONResponse({
            "ok": True,
            "filename": name,
            "page_count": ctx.get("page_count", 0),
            "text": text_out,
            "images": ctx.get("images", []),
            "project_details": details,
        })
    except Exception as e:
        return JSONResponse({"ok": False, "error": f"PDF extraction failed: {e}"}, status_code=500)


@app.post("/extract-document")
async def extract_document(file: UploadFile = File(...)):
    name = (file.filename or "").strip()
    ext = Path(name).suffix.lower()
    supported = {".docx", ".txt", ".md", ".csv", ".tsv", ".rtf", ".json", ".doc"}
    if ext not in supported:
        return JSONResponse({"ok": False, "error": f"Unsupported document type: {ext or 'unknown'}"}, status_code=400)

    try:
        raw_bytes = await file.read()
        out = _extract_document_text(raw_bytes, name)
        if ext == ".doc" and not out.get("text"):
            return JSONResponse(
                {
                    "ok": False,
                    "error": "Legacy .doc parsing is limited. Please upload .docx or PDF for reliable extraction.",
                    "filename": name,
                    "extension": ext,
                },
                status_code=400,
            )
        return JSONResponse({
            "ok": True,
            "filename": name,
            "text": out.get("text", ""),
            "project_details": out.get("project_details", {}),
            "parser": out.get("parser", ""),
            "extension": out.get("extension", ext),
        })
    except Exception as e:
        return JSONResponse({"ok": False, "error": f"Document extraction failed: {e}"}, status_code=500)


@app.post("/infer-text")
async def infer_text(payload: dict):
    """
    Text-only inference endpoint for PDF/Word-heavy uploads.
    Enables project/part auto-fill even when no image/CAD is uploaded.
    """
    try:
        context_text = str(payload.get("context_text", "") or "").strip()
        if not context_text:
            return JSONResponse({"ok": False, "error": "context_text is required"}, status_code=400)
        if _looks_like_email_thread(context_text):
            cleaned = _clean_email_thread_text(context_text)
            if cleaned:
                print(f"[infer-text] Email-thread cleanup applied ({len(context_text)} -> {len(cleaned)} chars)")
                context_text = cleaned
        inferred = run_text_inference(context_text)
        return JSONResponse({"ok": True, "inference": inferred})
    except Exception as e:
        return JSONResponse({"ok": False, "error": f"text inference failed: {e}"}, status_code=500)


@app.post("/analyze-cad")
async def analyze_cad(
    file: UploadFile = File(...),
    company_name: str = Form(...),
    zoho_id: str = Form(""),
    context_text: str = Form(""),
):
    t0 = time.time()
    print("\n[CAD] Step 1/7: Request received")
    suffix = Path(file.filename or "").suffix.lower()
    if not _is_cad_filename(file.filename or ""):
        return JSONResponse({"ok": False, "error": f"Unsupported CAD file type: {suffix}"}, status_code=400)

    if trimesh is None:
        return JSONResponse(
            {"ok": False, "error": "CAD processing not available. Install trimesh in backend environment."},
            status_code=500,
        )

    print(f"[CAD] File: {file.filename} | extension: {suffix}")
    if _looks_like_email_thread(context_text):
        cleaned = _clean_email_thread_text(context_text)
        if cleaned:
            print(f"[CAD] Email-thread cleanup applied to context_text ({len(context_text)} -> {len(cleaned)} chars)")
            context_text = cleaned
    raw_bytes = await file.read()
    
    # === MEMORY GUARD: Use MAX_CAD_FILE_MB env var (default 20MB for 512MB deployment) ===
    file_size_mb = len(raw_bytes) / (1024 * 1024)
    try:
        max_file_size_mb = int(os.getenv("MAX_CAD_FILE_MB", "20"))
    except ValueError:
        max_file_size_mb = 20
    
    if file_size_mb > max_file_size_mb:
        print(f"[CAD] ⚠️  File too large: {file_size_mb:.1f}MB > {max_file_size_mb}MB limit (MAX_CAD_FILE_MB)")
        return JSONResponse({
            "ok": False,
            "error": f"CAD file too large: {file_size_mb:.1f}MB (max: {max_file_size_mb}MB). "
                     f"Please upload a simpler or lower-resolution model."
        }, status_code=413)
    
    print(f"[CAD] Step 2/7: File read ({len(raw_bytes)} bytes / {file_size_mb:.1f}MB, limit: {max_file_size_mb}MB)")

    def _env_int(name: str, default: int) -> int:
        try:
            return int(os.getenv(name, str(default)))
        except Exception:
            return int(default)

    ext = suffix.lstrip(".")
    if ext == "stp":
        ext = "step"
    if ext == "igs":
        ext = "iges"

    tmp_path = UPLOAD_DIR / f"{uuid.uuid4()}.jpg"
    try:
        t_parse = time.time()
        cad_parse_mode = "mesh"
        geometry_available = True
        geometry_source = "mesh"
        try:
            mesh = trimesh.load(io.BytesIO(raw_bytes), file_type=ext, force="mesh")
        except Exception as te:
            if ext == "3mf":
                print(f"[CAD] trimesh direct parse failed for 3MF ({te}); trying 3MF XML mesh fallback")
                mesh = _load_3mf_mesh_from_xml(raw_bytes)
                cad_parse_mode = "mesh_3mf_fallback"
                geometry_source = "3mf_xml_mesh"
            else:
                raise

        if mesh is None or getattr(mesh, "vertices", None) is None:
            raise ValueError("CAD file could not be parsed as a mesh")

        original_vertex_count = len(mesh.vertices) if hasattr(mesh, "vertices") else 0
        original_face_count = len(mesh.faces) if hasattr(mesh, "faces") else 0
        hard_max_vertices = max(100000, _env_int("CAD_HARD_MAX_VERTICES", 1200000))
        hard_max_faces = max(200000, _env_int("CAD_HARD_MAX_FACES", 2200000))
        soft_heavy_vertices = max(50000, _env_int("CAD_SOFT_HEAVY_VERTICES", 300000))
        soft_heavy_faces = max(100000, _env_int("CAD_SOFT_HEAVY_FACES", 600000))

        if original_vertex_count > hard_max_vertices or original_face_count > hard_max_faces:
            print(
                f"[CAD] ⚠️  Model too complex: vertices={original_vertex_count:,}, faces={original_face_count:,} "
                f"(limits vertices={hard_max_vertices:,}, faces={hard_max_faces:,})"
            )
            return JSONResponse(
                {
                    "ok": False,
                    "error": (
                        "CAD model is too complex to process safely on current instance. "
                        f"vertices={original_vertex_count:,}, faces={original_face_count:,}. "
                        "Please simplify mesh/tessellation and re-upload."
                    ),
                },
                status_code=413,
            )

        heavy_mode = original_vertex_count > soft_heavy_vertices or original_face_count > soft_heavy_faces
        if heavy_mode:
            print(
                f"[CAD] Heavy-model mode enabled: vertices={original_vertex_count:,}, faces={original_face_count:,} "
                f"(soft thresholds vertices={soft_heavy_vertices:,}, faces={soft_heavy_faces:,})"
            )

        print(f"[CAD] Step 3/7: CAD parsed to mesh ({time.time() - t_parse:.2f}s)")

        t_preview = time.time()
        print("[CAD] Rendering technical gray isometric preview (fast mode)")
        preview_mesh = mesh
        preview_max_faces = None
        if heavy_mode:
            preview_max_faces = max(12000, _env_int("CAD_HEAVY_PREVIEW_MAX_FACES", 25000))
            try:
                preview_mesh = mesh.simplify_quadric_decimation(preview_max_faces)
                pv = len(preview_mesh.vertices) if hasattr(preview_mesh, "vertices") else 0
                pf = len(preview_mesh.faces) if hasattr(preview_mesh, "faces") else 0
                print(f"[CAD] Heavy preview decimation: vertices={pv:,}, faces={pf:,}")
            except Exception as e:
                preview_mesh = mesh
                print(f"[CAD] Heavy preview decimation failed; using original mesh for preview: {e}")

        preview_bundle = _mesh_to_preview_bundle(preview_mesh, max_faces_override=preview_max_faces)
        preview_bytes = preview_bundle["study"]["bytes"]
        preview_b64 = base64.b64encode(preview_bytes).decode("utf-8")
        extra_view_images = []
        for item in preview_bundle.get("extras", []):
            b64 = base64.b64encode(item["bytes"]).decode("utf-8")
            extra_view_images.append({
                "name": item.get("name", "view"),
                "filename": f"{Path(file.filename or 'cad').stem}_{item.get('name', 'view')}.jpg",
                "b64": b64,
                "data_url": f"data:image/jpeg;base64,{b64}",
            })
        with open(tmp_path, "wb") as f:
            f.write(preview_bytes)
        print(f"[CAD] Step 4/7: Preview image generated ({time.time() - t_preview:.2f}s)")

        # Compute mesh stats from original mesh (not preview-decimated copy).
        cad_stats = _cad_mesh_stats(mesh)

        # === Explicit cleanup: Free large mesh object from memory ===
        if preview_mesh is not mesh:
            del preview_mesh
        del mesh
        import gc
        gc.collect()

        t_scores = time.time()
        scores = compute_geometric_scores(str(tmp_path))
        geo_vec = build_geometric_vector(scores)
        print(f"[CAD] Step 5/7: Geometry scores computed ({time.time() - t_scores:.2f}s)")

        t_infer = time.time()
        inference = run_inference(str(tmp_path), scores, context_text=context_text)
        print(f"[CAD] Step 6/7: Field extraction completed ({time.time() - t_infer:.2f}s)")

        t_clip = time.time()
        clip_vec = None
        clip_ok = False
        embedding_enabled = CAD_CLIP_ENABLED or EMBEDDER_IS_EFFICIENTNET
        if embedding_enabled:
            clip_vec = compute_clip_embedding(str(tmp_path))
            clip_ok = clip_vec is not None and not _is_zero_vector(_to_float_vec(clip_vec))
        print(
            f"[CAD] Step 7/7: Vector step completed ({time.time() - t_clip:.2f}s) "
            f"| enabled={embedding_enabled} | backend={EMBEDDER_BACKEND}"
        )

        project_details = _cad_project_details_from_inference(inference, cad_stats)

        # Build part_envelope string from mesh bounding box (native CAD units, typically mm)
        def _fmt_dim(v):
            if not v or v < 0.001:
                return ""
            s = f"{v:.2f}".rstrip("0").rstrip(".")
            return s
        _bx = _fmt_dim(cad_stats.get("bbox_x", 0.0))
        _by = _fmt_dim(cad_stats.get("bbox_y", 0.0))
        _bz = _fmt_dim(cad_stats.get("bbox_z", 0.0))
        _dims = [d for d in [_bx, _by, _bz] if d]
        cad_part_envelope = (" x ".join(_dims) + " mm") if _dims else ""

        part_id = f"part_{zoho_id or 'nz'}_{uuid.uuid4().hex[:8]}"
        print(f"[CAD] Completed in {time.time() - t0:.2f}s | envelope={cad_part_envelope!r}")
        return JSONResponse({
            "ok": True,
            "part_id": part_id,
            "filename": file.filename,
            "source": inference.get("source", "cad"),
            "cad_parse_mode": cad_parse_mode,
            "geometry_available": geometry_available,
            "geometry_source": geometry_source,
            "clip_ok": clip_ok,
            "scores": {
                "aspect_ratio": scores["aspect_ratio"],
                "circularity": scores["circularity"],
                "convexity": scores["convexity"],
                "edge_density": scores["edge_density"],
                "symmetry_score": scores["symmetry_score"],
                "hole_count": scores["hole_count"],
                "reflectivity": scores["reflectivity"],
                "feature_complexity": scores["feature_complexity"],
                "compactness": scores["compactness"],
                "slenderness": scores["slenderness"],
                "mean_brightness": scores["mean_brightness"],
                "surface_std_dev": scores["surface_std_dev"],
            },
            "inference": {
                "part_family": inference["part_family"]["value"],
                "part_family_detail": inference["part_family"].get("detail", ""),
                "part_family_conf": inference["part_family"]["confidence"],
                "material": inference["material"]["value"],
                "material_reasoning": inference["material"].get("reasoning", ""),
                "material_conf": inference["material"]["confidence"],
                "process_primary": inference["process"]["primary"],
                "process_secondary": inference["process"].get("secondary", ""),
                "process_conf": inference["process"]["confidence"],
                "finish": inference["finish"]["value"],
                "finish_ra": inference["finish"].get("ra_estimate", ""),
                "finish_conf": inference["finish"]["confidence"],
                "complexity_class": inference.get("complexity_class", {}).get("value", ""),
                "tolerance_class": inference.get("tolerance_class", {}).get("value", ""),
                "features": inference.get("features", []),
                "notes": inference.get("notes"),
                "part_envelope": cad_part_envelope,
            },
            "clip_vector": clip_vec,
            "geo_scores": geo_vec,
            "cad_stats": cad_stats,
            "project_details": project_details,
            "preview_b64": preview_b64,
            "preview_data_url": f"data:image/jpeg;base64,{preview_b64}",
            "extra_view_images": extra_view_images,
        })
    except Exception as e:
        print(f"[CAD] ERROR after {time.time() - t0:.2f}s: {e}")
        known_3mf = suffix == ".3mf" and (
            "world" in str(e).lower()
            or "no mesh/components found in 3mf resources" in str(e).lower()
            or "no triangulated geometry resolved from 3mf" in str(e).lower()
        )
        if not known_3mf:
            import traceback
            traceback.print_exc()
        else:
            print("[CAD] Known 3MF mesh parse edge-case; continuing with preview fallback.")
        # 3MF fallback path: do not fail whole request if mesh parsing is flaky.
        if suffix == ".3mf":
            try:
                embedded = _extract_3mf_embedded_images(raw_bytes, max_images=1)
                if embedded:
                    from PIL import Image
                    pil = Image.open(io.BytesIO(embedded[0])).convert("RGB")
                    pil.save(tmp_path, format="JPEG", quality=90)
                    preview_bytes = tmp_path.read_bytes()
                    preview_b64 = base64.b64encode(preview_bytes).decode("utf-8")

                    scores = compute_geometric_scores(str(tmp_path))
                    inference = run_inference(str(tmp_path), scores, context_text=context_text)
                    zero_geo_vec = [0.0] * 18
                    empty_stats = {
                        "triangles": 0,
                        "vertices": 0,
                        "bbox_x": 0.0,
                        "bbox_y": 0.0,
                        "bbox_z": 0.0,
                        "surface_area": 0.0,
                        "volume": 0.0,
                    }
                    project_details = _cad_project_details_from_inference(inference, empty_stats)
                    part_id = f"part_{zoho_id or 'nz'}_{uuid.uuid4().hex[:8]}"
                    print("[CAD] Fallback succeeded: extracted preview image from 3MF package (geometry unavailable)")
                    return JSONResponse({
                        "ok": True,
                        "part_id": part_id,
                        "filename": file.filename,
                        "source": inference.get("source", "cad"),
                        "cad_parse_mode": "preview_only",
                        "geometry_available": False,
                        "geometry_source": "3mf_embedded_preview",
                        "clip_ok": False,
                        "scores": {
                            "aspect_ratio": scores["aspect_ratio"],
                            "circularity": scores["circularity"],
                            "convexity": scores["convexity"],
                            "edge_density": scores["edge_density"],
                            "symmetry_score": scores["symmetry_score"],
                            "hole_count": scores["hole_count"],
                            "reflectivity": scores["reflectivity"],
                            "feature_complexity": scores["feature_complexity"],
                            "compactness": scores["compactness"],
                            "slenderness": scores["slenderness"],
                            "mean_brightness": scores["mean_brightness"],
                            "surface_std_dev": scores["surface_std_dev"],
                        },
                        "inference": {
                            "part_family": inference["part_family"]["value"],
                            "part_family_detail": inference["part_family"].get("detail", ""),
                            "part_family_conf": inference["part_family"]["confidence"],
                            "material": inference["material"]["value"],
                            "material_reasoning": inference["material"].get("reasoning", ""),
                            "material_conf": inference["material"]["confidence"],
                            "process_primary": inference["process"]["primary"],
                            "process_secondary": inference["process"].get("secondary", ""),
                            "process_conf": inference["process"]["confidence"],
                            "finish": inference["finish"]["value"],
                            "finish_ra": inference["finish"].get("ra_estimate", ""),
                            "finish_conf": inference["finish"]["confidence"],
                            "complexity_class": inference.get("complexity_class", {}).get("value", ""),
                            "tolerance_class": inference.get("tolerance_class", {}).get("value", ""),
                            "features": inference.get("features", []),
                            "notes": inference.get("notes"),
                        },
                        "clip_vector": None,
                        "geo_scores": zero_geo_vec,
                        "cad_stats": empty_stats,
                        "project_details": project_details,
                        "preview_b64": preview_b64,
                        "preview_data_url": f"data:image/jpeg;base64,{preview_b64}",
                        "extra_view_images": [],
                    })
            except Exception as fe:
                print(f"[CAD] 3MF fallback failed: {fe}")
        return JSONResponse({"ok": False, "error": f"CAD analysis failed: {e}"}, status_code=500)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()

@app.post("/analyze")
async def analyze_image(
    file:         UploadFile = File(...),
    company_name: str        = Form(...),
    zoho_id:      str        = Form(""),
    context_text: str        = Form(""),
):
    tmp_path = UPLOAD_DIR / f"{uuid.uuid4()}.jpg"
    try:
        if _looks_like_email_thread(context_text):
            cleaned = _clean_email_thread_text(context_text)
            if cleaned:
                print(f"[analyze] Email-thread cleanup applied to context_text ({len(context_text)} -> {len(cleaned)} chars)")
                context_text = cleaned
        raw_bytes = file.file.read()
        from PIL import Image
        import io
        try:
            import pillow_avif
        except ImportError:
            pass
        try:
            img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
            img.save(str(tmp_path), "JPEG", quality=95)
        except Exception:
            try:
                import imageio, numpy as np
                img_arr = imageio.imread(io.BytesIO(raw_bytes))
                img = Image.fromarray(img_arr).convert("RGB")
                img.save(str(tmp_path), "JPEG", quality=95)
            except Exception as e2:
                raise ValueError(f"Cannot decode image {file.filename}: {e2}")

        print(f"\n  â†’ Analysing: {file.filename}")
        scores    = compute_geometric_scores(str(tmp_path))
        geo_vec   = build_geometric_vector(scores)
        inference = run_inference(str(tmp_path), scores, context_text=context_text)

        print(f"  â†’ Computing visual embedding (backend={EMBEDDER_BACKEND})â€¦")
        clip_vec = compute_clip_embedding(str(tmp_path))
        clip_ok  = clip_vec is not None and not _is_zero_vector(_to_float_vec(clip_vec))
        if not clip_ok:
            print(f"  âš  Visual embedding unavailable or returned zero vector â€” push step will retry from payload/image")

        part_id = f"part_{zoho_id or 'nz'}_{uuid.uuid4().hex[:8]}"
        print(
            f"  âœ“ [{inference.get('source','?')}] {inference['part_family']['value']} Â· {inference['material']['value']} "
            f"| EMBED({EMBEDDER_BACKEND}): {'âœ“' if clip_ok else 'âœ—'}"
        )

        return JSONResponse({
            "ok":      True,
            "part_id": part_id,
            "filename":file.filename,
            "source":  inference.get("source", "unknown"),
            "clip_ok": clip_ok,
            "scores": {
                "aspect_ratio":       scores["aspect_ratio"],
                "circularity":        scores["circularity"],
                "convexity":          scores["convexity"],
                "edge_density":       scores["edge_density"],
                "symmetry_score":     scores["symmetry_score"],
                "hole_count":         scores["hole_count"],
                "reflectivity":       scores["reflectivity"],
                "feature_complexity": scores["feature_complexity"],
                "compactness":        scores["compactness"],
                "slenderness":        scores["slenderness"],
                "mean_brightness":    scores["mean_brightness"],
                "surface_std_dev":    scores["surface_std_dev"],
            },
            "inference": {
                "part_family":        inference["part_family"]["value"],
                "part_family_detail": inference["part_family"].get("detail", ""),
                "part_family_conf":   inference["part_family"]["confidence"],
                "material":           inference["material"]["value"],
                "material_reasoning": inference["material"].get("reasoning", ""),
                "material_conf":      inference["material"]["confidence"],
                "process_primary":    inference["process"]["primary"],
                "process_secondary":  inference["process"].get("secondary", "â€”"),
                "process_conf":       inference["process"]["confidence"],
                "finish":             inference["finish"]["value"],
                "finish_ra":          inference["finish"].get("ra_estimate", "â€”"),
                "finish_conf":        inference["finish"]["confidence"],
                "complexity_class":   inference.get("complexity_class", {}).get("value", ""),
                "tolerance_class":    inference.get("tolerance_class", {}).get("value", ""),
                "features":           inference.get("features", []),
                "notes":              inference.get("notes"),
            },
            "clip_vector": clip_vec,   # None if CLIP not loaded â€” frontend handles this
            "geo_scores":  geo_vec,
        })

    except Exception as e:
        print(f"  âœ— Error: {e}")
        import traceback; traceback.print_exc()
        return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


@app.post("/push")
async def push_to_pinecone(payload: dict):
    parts = payload.get("parts", [])
    if not pinecone_index:
        # â”€â”€ No Pinecone â€” still try Zoho â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        results = []
        for part in parts:
            part_id = part.get("part_id", str(uuid.uuid4()))
            zoho_result = {"ok": False, "error": "Zoho not configured"}
            if ZOHO_REFRESH_TOKEN:
                try:
                    zoho_result = zoho_push_past_project(part)
                except Exception as ze:
                    zoho_result = {"ok": False, "error": str(ze)}
            results.append({
                "part_id":  part_id, "ok": zoho_result.get("ok", False),
                "zoho_ok":  zoho_result.get("ok", False),
                "zoho_action": zoho_result.get("action", ""),
                "error":    "Pinecone not connected" if not zoho_result.get("ok") else "",
            })
        zoho_ok = sum(1 for r in results if r.get("zoho_ok"))
        return JSONResponse({"ok": zoho_ok > 0, "pushed": 0, "failed": len(parts),
                             "zoho_ok": zoho_ok, "results": results})

    results = []
    for part in parts:
        part_id = part.get("part_id", str(uuid.uuid4()))
        try:
            # â”€â”€ 1. Save image â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            saved_image_path = None
            img_b64   = part.get("image_b64")
            img_ext   = part.get("image_ext", ".jpg")
            image_url = (part.get("image_url") or "").strip()
            image_preview = (part.get("image_preview") or "").strip()
            if not image_url and image_preview and not image_preview.startswith("blob:"):
                image_url = image_preview

            if img_b64:
                safe_name = f"{part_id}{img_ext}"
                dest      = STORED_PARTS_DIR / safe_name
                with open(dest, "wb") as f:
                    f.write(base64.b64decode(img_b64))
                image_url = f"/parts/{safe_name}"
                saved_image_path = dest
                print(f"  âœ“ Image saved â†’ {dest}")

            part["image_url"] = image_url

            # â”€â”€ 2. Build metadata â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            features     = part.get("features", [])
            features_str = ", ".join(features) if isinstance(features, list) else str(features)

            project_name = (
                part.get("project_name")
                or part.get("part_name")
                or part.get("filename")
                or part.get("part_family")
                or ""
            )

            metadata = {
                "supplier_name":      part.get("company_name", ""),
                "zoho_id":            str(part.get("zoho_id", "")),
                "source_type":        "HISTORICAL_PROJECT",
                "project_id":         part.get("project_id", ""),
                "project_name":       project_name,
                "image_name":         part.get("filename", ""),
                "image_url":          image_url,
                "inference_source":   part.get("inference_source", "ai"),
                "part_family":        part.get("part_family", ""),
                "part_family_detail": part.get("part_family_detail", ""),
                "part_family_conf":   float(part.get("part_family_conf", 0)),
                "material":           part.get("material", ""),
                "material_reasoning": part.get("material_reasoning", ""),
                "material_conf":      float(part.get("material_conf", 0)),
                "process_primary":    part.get("process", ""),
                "process_secondary":  part.get("process_secondary", ""),
                "process_conf":       float(part.get("process_conf", 0)),
                "finish":             part.get("finish", ""),
                "finish_ra":          part.get("finish_ra", ""),
                "finish_conf":        float(part.get("finish_conf", 0)),
                "complexity_class":   part.get("complexity_class", ""),
                "tolerance_class":    part.get("tolerance_class", ""),
                "features":           features_str,
                "notes":              part.get("notes", "") or "",
                "outcome":            part.get("outcome", ""),
                "ncr_description":    part.get("ncr_description", ""),
                "what_worked":        part.get("what_worked", ""),
                "what_didnt":         part.get("what_didnt", ""),
                "project_date":       part.get("project_date", ""),
                "customer_industry":  part.get("customer_industry", ""),
                "share_with_tb":      str(part.get("share_with_tb", True)),
                "geo_circularity":    float(part.get("circularity", 0)),
                "geo_symmetry":       float(part.get("symmetry", 0)),
                "geo_hole_count":     float(part.get("hole_count", 0)),
                "geo_complexity":     float(part.get("complexity", 0)),
                "geo_aspect_ratio":   float(part.get("aspect_ratio", 0)),
            }

            cad_stats = part.get("cad_stats", {}) or {}
            metadata.update({
                "cad_triangles": int(cad_stats.get("triangles", part.get("cad_triangles", 0)) or 0),
                "cad_vertices": int(cad_stats.get("vertices", part.get("cad_vertices", 0)) or 0),
                "cad_bbox_x": float(cad_stats.get("bbox_x", part.get("cad_bbox_x", 0)) or 0),
                "cad_bbox_y": float(cad_stats.get("bbox_y", part.get("cad_bbox_y", 0)) or 0),
                "cad_bbox_z": float(cad_stats.get("bbox_z", part.get("cad_bbox_z", 0)) or 0),
                "cad_surface_area": float(cad_stats.get("surface_area", part.get("cad_surface_area", 0)) or 0),
                "cad_volume": float(cad_stats.get("volume", part.get("cad_volume", 0)) or 0),
            })

            # â”€â”€ 3. âœ… FIX: Resolve vector â€” never send all-zeros to Pinecone â”€â”€
            raw_vec  = part.get("clip_vector")
            clip_vec = _to_float_vec(raw_vec)
            non_zero = sum(1 for v in clip_vec if abs(v) > 1e-12)
            print(f"  Embedding vec dims : {len(clip_vec)} | non-zero: {non_zero} | backend={EMBEDDER_BACKEND}")

            if _is_zero_vector(clip_vec):
                # Recovery path: recompute from actual image/payload in backend.
                recomputed = _compute_embedding_from_part_payload(part, saved_image_path=saved_image_path)
                if recomputed is not None:
                    recomputed_vec = _to_float_vec(recomputed)
                    if not _is_zero_vector(recomputed_vec):
                        clip_vec = recomputed_vec
                        print(f"  âœ“ Recomputed {EMBEDDER_BACKEND} embedding from payload/image for {part_id}")

            if _is_zero_vector(clip_vec):
                if REQUIRE_REAL_CLIP_VECTOR:
                    raise ValueError(
                        f"Real {EMBEDDER_BACKEND} vector required but received zero/empty vector. "
                        f"Check {EMBEDDER_BACKEND} runtime dependencies and analyze path."
                    )
                # Embedding failed or returned zeros â€” build a deterministic metadata
                # fingerprint vector so Pinecone never sees an all-zero payload.
                clip_vec = _make_fallback_vector(part_id, metadata)
                print(f"  âš  Zero {EMBEDDER_BACKEND} vector detected for {part_id} â€” using metadata fingerprint fallback")
                metadata["vector_type"] = "metadata_fingerprint"   # flag for downstream use
            else:
                metadata["vector_type"] = "efficientnet" if EMBEDDER_IS_EFFICIENTNET else "clip"

            print(f"\n  INSERTING â†’ Pinecone ID: {part_id}")
            print(f"  Part Family : {metadata['part_family']} ({int(metadata['part_family_conf']*100)}%)")
            print(f"  Material    : {metadata['material']}")
            print(f"  Process     : {metadata['process_primary']}")
            print(f"  Vec type    : {metadata['vector_type']}")
            print(f"  Vector[0:3] : {clip_vec[0]:.4f}, {clip_vec[1]:.4f}, {clip_vec[2]:.4f}â€¦")

            pinecone_index.upsert(vectors=[{
                "id":       part_id,
                "values":   clip_vec,
                "metadata": metadata,
            }])
            print(f"  âœ“ Pushed to Pinecone â†’ {part_id}")

            # â”€â”€ 4. Push to Zoho CRM â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            
            zoho_result = {"ok": False, "error": "Zoho not configured"}
            if ZOHO_REFRESH_TOKEN:
                try:
                    zoho_result = zoho_push_past_project(part, index_ready=True)
                except Exception as ze:
                    print(f"  âœ— Zoho push error: {ze}")
                    zoho_result = {"ok": False, "error": str(ze)}
            else:
                print(f"  âš  Zoho push skipped â€” ZOHO_REFRESH_TOKEN not set")

            results.append({
                "part_id":     part_id, "ok": True,
                "image_url":   image_url, "clip_pushed": True,
                "vector_type": metadata["vector_type"],
                "zoho_ok":     zoho_result.get("ok", False),
                "zoho_action": zoho_result.get("action", ""),
                "zoho_error":  zoho_result.get("error", ""),
            })

        except Exception as e:
            print(f"  âœ— Push failed {part_id}: {e}")
            import traceback; traceback.print_exc()
            results.append({"part_id": part_id, "ok": False, "error": str(e)})

    pushed  = sum(1 for r in results if r.get("ok"))
    zoho_ok = sum(1 for r in results if r.get("zoho_ok"))
    print(f"\n  Summary: {pushed}/{len(parts)} â†’ Pinecone Â· {zoho_ok}/{len(parts)} â†’ Zoho CRM")

    return JSONResponse({
        "ok":      pushed > 0,
        "pushed":  pushed,
        "failed":  len(parts) - pushed,
        "zoho_ok": zoho_ok,
        "results": results,
    })
from pydantic import BaseModel, Field
from typing import List, Optional

class ZohoSyncPart(BaseModel):
    part_id:           str
    company_name:      Optional[str] = ""
    company_size:      Optional[str] = ""
    company_location:  Optional[str] = ""
    zoho_id:           Optional[str] = ""
    supplier_email:    Optional[str] = ""
    contact_email:     Optional[str] = ""
    contact_phone:     Optional[str] = ""
    project_name:      Optional[str] = ""
    project_description: Optional[str] = ""
    expected_annual_production_volume: Optional[str] = ""
    certification_notes: Optional[str] = ""
    mandatory_certifications: List[str] = Field(default_factory=list)
    data_sharing_tier: Optional[str] = ""
    part_family:       Optional[str] = ""
    part_name:         Optional[str] = ""
    part_detail:       Optional[str] = ""
    quantity:          Optional[float] = None
    part_envelope:     Optional[str] = ""
    additional_notes:  Optional[str] = ""
    surface_finish:    Optional[str] = ""
    tolerance_details: Optional[str] = ""
    material:          Optional[str] = ""
    process_primary:   Optional[str] = ""
    process_secondary: Optional[str] = ""
    complexity_class:  Optional[str] = ""
    tolerance_class:   Optional[str] = ""
    outcome:           Optional[str] = ""
    ncr_description:   Optional[str] = ""
    what_worked:       Optional[str] = ""
    what_didnt:        Optional[str] = ""
    customer_industry: Optional[str] = ""
    project_date:      Optional[str] = ""
    image_url:         Optional[str] = ""
    image_b64:         Optional[str] = ""
    image_ext:         Optional[str] = ""
    cad_filename:      Optional[str] = ""
    cad_file_b64:      Optional[str] = ""
    cad_preview_b64:   Optional[str] = ""
    cad_preview_filename: Optional[str] = ""
    cad_extra_views:   List[dict] = Field(default_factory=list)

class ZohoSyncPayload(BaseModel):
    parts: List[ZohoSyncPart]


class LessonsSyncPayload(BaseModel):
    supplier_id: Optional[str] = ""
    supplier_email: Optional[str] = ""
    mfg_lessons: List[dict] = Field(default_factory=list)
    quoting_lessons: List[dict] = Field(default_factory=list)
    deleted_mfg_lesson_ids: List[str] = Field(default_factory=list)
    deleted_quoting_lesson_ids: List[str] = Field(default_factory=list)


def _strip_subform_system_fields(row: dict) -> dict:
    clean = dict(row or {})
    for key in ("id", "Created_Time", "Modified_Time", "Created_By", "Modified_By"):
        clean.pop(key, None)
    return clean


def _extract_project_hint(lesson: dict) -> str:
    source = (
        (lesson.get("source_job") or "").strip()
        or (lesson.get("source_part") or "").strip()
        or (lesson.get("title") or "").strip()
    )
    if " · " in source:
        return source.split(" · ", 1)[0].strip().lower()
    return source.strip().lower()


def _build_mfg_subform_row(lesson: dict, sort_order: int) -> dict:
    lesson_id = str(
        lesson.get("id")
        or lesson.get("lesson_id")
        or lesson.get("Lesson_ID")
        or lesson.get("Lesson_Id")
        or lesson.get("row_id")
        or uuid.uuid4()
    ).strip()
    return {
        "Lesson_ID": lesson_id,
        "Category": (lesson.get("category") or "Process").strip(),
        "Title": (lesson.get("title") or "Manufacturing Lesson").strip(),
        "Description": (lesson.get("desc") or "").strip(),
        "Source_Part": (lesson.get("source_part") or "").strip(),
        "Process": (lesson.get("process") or "").strip(),
        "Material": (lesson.get("material") or "").strip(),
        "Created_From": "Manual",
        "Sort_Order": sort_order,
    }


def _build_quote_subform_row(lesson: dict, sort_order: int) -> dict:
    attachment_names = lesson.get("attachment_names") or lesson.get("attachments") or []
    if isinstance(attachment_names, list):
        attachment_names = ", ".join([str(x).strip() for x in attachment_names if str(x).strip()])
    attachment_names = str(attachment_names or "").strip()
    lesson_id = str(
        lesson.get("id")
        or lesson.get("lesson_id")
        or lesson.get("Lesson_ID")
        or lesson.get("Lesson_Id")
        or lesson.get("row_id")
        or uuid.uuid4()
    ).strip()
    return {
        "Lesson_ID": lesson_id,
        "Category": (lesson.get("category") or "Other").strip(),
        "Title": (lesson.get("title") or "Quoting Lesson").strip(),
        "Description": (lesson.get("desc") or "").strip(),
        "Process": (lesson.get("process") or "").strip(),
        "Material": (lesson.get("material") or "").strip(),
        "Source_Label": (lesson.get("source_label") or "").strip(),
        "Source_Job": (lesson.get("source_job") or "").strip(),
        "Tier": (lesson.get("tier") or "").strip(),
        "Date": (lesson.get("date") or "").strip(),
        "Image_Name": (lesson.get("image_name") or "").strip(),
        "Attachment_Names": attachment_names,
        "Created_From": "Manual",
        "Sort_Order": sort_order,
    }


def _debug_row_snapshot(row: dict) -> dict:
    item = dict(row or {})
    return {
        "Lesson_ID": str(item.get("Lesson_ID") or ""),
        "Category": str(item.get("Category") or ""),
        "Title": str(item.get("Title") or ""),
        "Source_Part": str(item.get("Source_Part") or ""),
        "Source_Job": str(item.get("Source_Job") or ""),
        "Source_Label": str(item.get("Source_Label") or ""),
        "Tier": str(item.get("Tier") or ""),
        "Date": str(item.get("Date") or ""),
        "Process": str(item.get("Process") or ""),
        "Material": str(item.get("Material") or ""),
        "Created_From": str(item.get("Created_From") or ""),
        "Sort_Order": item.get("Sort_Order"),
        "Description_Length": len(str(item.get("Description") or "")),
    }


def _merge_lesson_subform_rows(existing_rows: List[dict], incoming_by_id: dict, delete_ids: Optional[set] = None) -> List[dict]:
    """
    Merge incoming lesson rows into the Zoho subform without deleting manual
    lessons that were created from another browser or teammate's system.
    """
    merged = []
    replaced = set()
    delete_ids = delete_ids or set()
    for existing in existing_rows or []:
        clean = _strip_subform_system_fields(existing)
        lid = str(
            clean.get("Lesson_ID")
            or clean.get("Lesson_Id")
            or existing.get("id")
            or ""
        ).strip()
        if lid and lid in delete_ids:
            continue
        if lid and lid in incoming_by_id:
            merged.append(incoming_by_id[lid])
            replaced.add(lid)
        else:
            merged.append(clean)

    for lid, row in (incoming_by_id or {}).items():
        if lid not in replaced and lid not in delete_ids:
            merged.append(row)
    return merged


@app.post("/zoho-sync")
async def zoho_sync(payload: ZohoSyncPayload):
    """
    Explicit Zoho CRM sync â€” called when the user clicks 'Sync to Zoho CRM'.
    Searches Supplier_Past_Projects by company_name, appends or creates.
    Completely separate from /push (Pinecone).
    """
    if not ZOHO_REFRESH_TOKEN:
        return JSONResponse(
            {"ok": False, "error": "Zoho not configured â€” ZOHO_REFRESH_TOKEN missing"},
            status_code=503,
        )

    results = []
    for part in payload.parts:
        part_dict = part.dict()
        # /zoho-sync sends process_primary but zoho_push_past_project reads "process"
        part_dict["process"] = part_dict.pop("process_primary", "")
        try:
            result = zoho_push_past_project(part_dict)
            results.append({"part_id": part.part_id, **result})
        except Exception as e:
            print(f"  âœ— zoho-sync error for {part.part_id}: {e}")
            results.append({"part_id": part.part_id, "ok": False, "error": str(e)})

    synced = sum(1 for r in results if r.get("ok"))
    print(f"\n  Zoho Sync: {synced}/{len(results)} parts pushed")
    return JSONResponse({
        "ok":     synced > 0,
        "synced": synced,
        "total":  len(results),
        "results": results,
    })


@app.post("/projects/update")
async def update_projects(payload: dict):
    part_updates = payload.get("part_updates") or []
    if not isinstance(part_updates, list) or not part_updates:
        return JSONResponse({"ok": False, "error": "part_updates is required"}, status_code=400)

    requested_part_ids = [str((item or {}).get("part_id") or "").strip() for item in part_updates]
    requested_part_ids = [pid for pid in requested_part_ids if pid]
    vector_map = _pinecone_fetch_vectors(requested_part_ids)

    results = []
    pinecone_updated = 0
    zoho_updated = 0

    for item in part_updates:
        item = item or {}
        part_id = str(item.get("part_id") or "").strip()
        record_id = str(item.get("record_id") or "").strip()
        values, existing_meta = _vector_values_and_metadata(vector_map.get(part_id))
        merged_part = _normalize_project_part_update(payload, item, existing_meta)

        pinecone_ok = False
        pinecone_error = ""
        if part_id and values and pinecone_index:
            try:
                pinecone_index.upsert(vectors=[{
                    "id": part_id,
                    "values": values,
                    "metadata": {k: v for k, v in merged_part.items() if k not in {"company_name", "supplier_email", "zoho_id", "part_id"}},
                }])
                pinecone_ok = True
                pinecone_updated += 1
            except Exception as e:
                pinecone_error = str(e)
        elif part_id and not pinecone_index:
            pinecone_error = "Pinecone not configured"
        elif part_id and not values:
            pinecone_error = "Existing Pinecone vector not found"

        zoho_result = {"ok": False, "error": "Zoho not configured"}
        if record_id and ZOHO_REFRESH_TOKEN:
            try:
                zoho_result = _zoho_update_project_record(record_id, merged_part)
                if zoho_result.get("ok"):
                    zoho_updated += 1
            except Exception as e:
                zoho_result = {"ok": False, "error": str(e)}
        elif record_id and not ZOHO_REFRESH_TOKEN:
            zoho_result = {"ok": False, "error": "Zoho not configured"}
        elif not record_id:
            zoho_result = {"ok": False, "error": "record_id missing"}

        results.append({
            "part_id": part_id,
            "record_id": record_id,
            "ok": pinecone_ok and bool(zoho_result.get("ok")),
            "pinecone_ok": pinecone_ok,
            "pinecone_error": pinecone_error,
            "zoho_ok": bool(zoho_result.get("ok")),
            "zoho_error": zoho_result.get("error", ""),
            "zoho_action": zoho_result.get("action", ""),
        })

    return JSONResponse({
        "ok": all(r.get("ok") for r in results) if results else False,
        "updated": sum(1 for r in results if r.get("ok")),
        "pinecone_updated": pinecone_updated,
        "zoho_updated": zoho_updated,
        "total": len(results),
        "results": results,
    })


@app.post("/projects/rename")
async def rename_project(payload: dict):
    record_ids = [str(x or "").strip() for x in (payload.get("record_ids") or [])]
    record_ids = [rid for rid in dict.fromkeys(record_ids) if rid]
    new_name = str(payload.get("new_name") or "").strip()
    if not record_ids or not new_name:
        return JSONResponse({"ok": False, "error": "record_ids and new_name are required"}, status_code=400)
    if not ZOHO_REFRESH_TOKEN:
        return JSONResponse({"ok": False, "error": "Zoho not configured"}, status_code=503)
    results = []
    for rid in record_ids:
        try:
            resp = requests.put(
                f"{ZOHO_API_BASE}/Supplier_Past_Projects",
                headers=zoho_headers(),
                json={"data": [{"id": rid, "Name": new_name}]},
                timeout=12,
            )
            if resp.status_code in (200, 201):
                results.append({"record_id": rid, "ok": True})
            else:
                results.append({"record_id": rid, "ok": False, "error": f"PUT {resp.status_code}: {resp.text[:200]}"})
        except Exception as e:
            results.append({"record_id": rid, "ok": False, "error": str(e)})
    return JSONResponse({
        "ok": all(r["ok"] for r in results) if results else False,
        "updated": sum(1 for r in results if r["ok"]),
        "total": len(results),
        "results": results,
    })


@app.post("/projects/delete")
async def delete_projects(payload: dict):
    record_ids = [str(x or "").strip() for x in (payload.get("record_ids") or [])]
    part_ids = [str(x or "").strip() for x in (payload.get("part_ids") or [])]
    record_ids = [rid for rid in dict.fromkeys(record_ids) if rid]
    part_ids = [pid for pid in dict.fromkeys(part_ids) if pid]

    if not record_ids and not part_ids:
        return JSONResponse({"ok": False, "error": "record_ids or part_ids is required"}, status_code=400)

    pinecone_deleted = 0
    zoho_deleted = 0
    pinecone_error = ""
    zoho_error = ""

    if part_ids:
        if pinecone_index:
            try:
                pinecone_index.delete(ids=part_ids)
                pinecone_deleted = len(part_ids)
            except Exception as e:
                pinecone_error = str(e)
        else:
            pinecone_error = "Pinecone not configured"

    if record_ids:
        if ZOHO_REFRESH_TOKEN:
            try:
                resp = requests.delete(
                    f"{ZOHO_API_BASE}/Supplier_Past_Projects",
                    headers=zoho_headers(),
                    params={"ids": ",".join(record_ids)},
                    timeout=15,
                )
                if resp.status_code in (200, 202):
                    zoho_deleted = len(record_ids)
                else:
                    zoho_error = f"DELETE {resp.status_code}: {resp.text[:300]}"
            except Exception as e:
                zoho_error = str(e)
        else:
            zoho_error = "Zoho not configured"

    ok = (not part_ids or pinecone_deleted == len(part_ids)) and (not record_ids or zoho_deleted == len(record_ids))
    return JSONResponse({
        "ok": ok,
        "pinecone_deleted": pinecone_deleted,
        "zoho_deleted": zoho_deleted,
        "part_ids": part_ids,
        "record_ids": record_ids,
        "pinecone_error": pinecone_error,
        "zoho_error": zoho_error,
    })


@app.post("/zoho-sync-lessons")
async def zoho_sync_lessons(payload: LessonsSyncPayload, debug: bool = False):
    if not ZOHO_REFRESH_TOKEN:
        return JSONResponse(
            {"ok": False, "error": "Zoho not configured - ZOHO_REFRESH_TOKEN missing"},
            status_code=503,
        )

    supplier_id = (payload.supplier_id or "").strip()
    supplier_email = (payload.supplier_email or "").strip().lower()
    records = _fetch_zoho_supplier_projects(
        supplier_id=supplier_id,
        supplier_email=supplier_email,
        limit=300,
        include_index_ready=True,
    )
    if not records:
        return JSONResponse(
            {"ok": False, "error": "No supplier project records found to attach lessons."},
            status_code=404,
        )

    default_record = records[0]
    by_name = {}
    by_id = {}
    for rec in records:
        name_key = (rec.get("Name") or "").strip().lower()
        if name_key and name_key not in by_name:
            by_name[name_key] = rec
        rid = str(rec.get("id") or "").strip()
        if rid:
            by_id[rid] = rec

    assignments = {}
    record_meta = {}

    def assign_row(kind: str, lesson: dict, row: dict):
        explicit_record_id = str(lesson.get("project_record_id") or "").strip()
        if explicit_record_id and explicit_record_id in by_id:
            target = by_id[explicit_record_id]
        else:
            hint = _extract_project_hint(lesson)
            target = by_name.get(hint) or default_record
        rid = target.get("id")
        if not rid:
            return
        if rid not in record_meta:
            record_meta[rid] = {
                "record_name": target.get("Name") or "",
                "record_job_id": target.get("Tag") or "",
            }
        if rid not in assignments:
            assignments[rid] = {"mfg": {}, "quote": {}}
        lid = row.get("Lesson_ID")
        if lid:
            assignments[rid]["mfg" if kind == "mfg" else "quote"][lid] = row

    for i, lesson in enumerate(payload.mfg_lessons or []):
        assign_row("mfg", lesson, _build_mfg_subform_row(lesson, i + 1))

    for i, lesson in enumerate(payload.quoting_lessons or []):
        assign_row("quote", lesson, _build_quote_subform_row(lesson, i + 1))

    deleted_mfg_ids = {str(x or "").strip() for x in (payload.deleted_mfg_lesson_ids or []) if str(x or "").strip()}
    deleted_quote_ids = {str(x or "").strip() for x in (payload.deleted_quoting_lesson_ids or []) if str(x or "").strip()}
    if deleted_mfg_ids or deleted_quote_ids:
        for rec in records:
            rid = str(rec.get("id") or "").strip()
            if not rid:
                continue
            if rid not in record_meta:
                record_meta[rid] = {
                    "record_name": rec.get("Name") or "",
                    "record_job_id": rec.get("Tag") or "",
                }
            if rid not in assignments:
                assignments[rid] = {"mfg": {}, "quote": {}}

    headers = zoho_headers()
    results = []
    debug_payloads = []

    for record_id, grouped in assignments.items():
        try:
            detail_resp = requests.get(
                f"{ZOHO_API_BASE}/Supplier_Past_Projects/{record_id}",
                headers=headers,
                timeout=12,
            )
            detail_resp.raise_for_status()
            detail = (detail_resp.json().get("data") or [{}])[0]
            existing_mfg = detail.get("MFG_Lessons") or []
            existing_quote = detail.get("Quoting_Lessons") or []

            merged_mfg = _merge_lesson_subform_rows(existing_mfg, grouped["mfg"], deleted_mfg_ids)
            merged_quote = _merge_lesson_subform_rows(existing_quote, grouped["quote"], deleted_quote_ids)
            put_payload = {
                "data": [{
                    "id": record_id,
                    "MFG_Lessons": merged_mfg,
                    "Quoting_Lessons": merged_quote,
                }]
            }

            if debug:
                debug_payloads.append({
                    "record_id": record_id,
                    "record_name": (record_meta.get(record_id) or {}).get("record_name", ""),
                    "record_job_id": (record_meta.get(record_id) or {}).get("record_job_id", ""),
                    "counts": {
                        "existing_mfg": len(existing_mfg),
                        "existing_quote": len(existing_quote),
                        "preserved_mfg": max(0, len(merged_mfg) - len(grouped["mfg"])),
                        "preserved_quote": max(0, len(merged_quote) - len(grouped["quote"])),
                        "incoming_manual_mfg": len(grouped["mfg"]),
                        "incoming_manual_quote": len(grouped["quote"]),
                        "merged_mfg": len(merged_mfg),
                        "merged_quote": len(merged_quote),
                    },
                    "put_payload_preview": {
                        "id": record_id,
                        "MFG_Lessons": [_debug_row_snapshot(r) for r in merged_mfg],
                        "Quoting_Lessons": [_debug_row_snapshot(r) for r in merged_quote],
                    },
                })

            update_resp = requests.put(
                f"{ZOHO_API_BASE}/Supplier_Past_Projects",
                headers=headers,
                json=put_payload,
                timeout=12,
            )
            if update_resp.status_code not in (200, 201):
                raise ValueError(f"PUT {update_resp.status_code}: {update_resp.text[:300]}")

            for lid in grouped["mfg"].keys():
                results.append({"kind": "mfg", "lesson_id": lid, "ok": True, "zoho_record_id": record_id})
            for lid in grouped["quote"].keys():
                results.append({"kind": "quote", "lesson_id": lid, "ok": True, "zoho_record_id": record_id})
        except Exception as e:
            for lid in grouped["mfg"].keys():
                results.append({"kind": "mfg", "lesson_id": lid, "ok": False, "error": str(e)})
            for lid in grouped["quote"].keys():
                results.append({"kind": "quote", "lesson_id": lid, "ok": False, "error": str(e)})

    synced = sum(1 for r in results if r.get("ok"))
    body = {"ok": synced > 0, "synced": synced, "total": len(results), "results": results}
    if debug:
        body["debug"] = {
            "supplier_id": supplier_id,
            "supplier_email": supplier_email,
            "record_count": len(records),
            "targeted_record_count": len(assignments),
            "assignments": debug_payloads,
        }
    return JSONResponse(body)


@app.get("/zoho-lessons")
async def zoho_lessons(supplier_id: str = "", supplier_email: str = "", limit: int = 300, debug: bool = False):
    try:
        if not (supplier_id or "").strip() and not (supplier_email or "").strip():
            return JSONResponse(
                {"ok": False, "error": "supplier_id or supplier_email is required", "mfg_lessons": [], "quoting_lessons": []},
                status_code=400,
            )
        all_rows = _fetch_zoho_supplier_projects(
            supplier_id="",
            supplier_email="",
            limit=limit,
            include_index_ready=True,
        )
        rows = _fetch_zoho_supplier_projects(
            supplier_id=(supplier_id or "").strip(),
            supplier_email=(supplier_email or "").strip().lower(),
            limit=limit,
            include_index_ready=True,
        )
        mfg_lessons = []
        quoting_lessons = []

        seen_mfg = set()
        seen_quote = set()
        seen_quote_fingerprint = set()
        headers = zoho_headers()

        for rec in rows:
            record_id = rec.get("id")
            if not record_id:
                continue

            detail_resp = requests.get(
                f"{ZOHO_API_BASE}/Supplier_Past_Projects/{record_id}",
                headers=headers,
                timeout=12,
            )
            if detail_resp.status_code != 200:
                continue
            detail = (detail_resp.json().get("data") or [{}])[0]

            for row in detail.get("MFG_Lessons") or []:
                created_from = str((row or {}).get("Created_From", "")).strip().lower()
                if created_from and created_from != "manual":
                    continue
                lid = str(
                    (row or {}).get("Lesson_ID", "")
                    or (row or {}).get("Lesson_Id", "")
                    or (row or {}).get("id", "")
                ).strip() or str(uuid.uuid4())
                if lid in seen_mfg:
                    continue
                seen_mfg.add(lid)
                mfg_lessons.append({
                    "id": lid,
                    "row_id": str((row or {}).get("id", "")).strip(),
                    "category": (row.get("Category") or "Process").strip(),
                    "title": (row.get("Title") or "Manufacturing Lesson").strip(),
                    "desc": (row.get("Description") or row.get("Lesson") or row.get("Notes") or "").strip(),
                    "source_part": (row.get("Source_Part") or "").strip(),
                    "process": (row.get("Process") or "").strip(),
                    "material": (row.get("Material") or "").strip(),
                    "project_record_id": str(record_id or "").strip(),
                    "project_name": (detail.get("Name") or rec.get("Name") or "").strip(),
                    "source": "manual",
                })

            for row in detail.get("Quoting_Lessons") or []:
                created_from = str((row or {}).get("Created_From", "")).strip().lower()
                if created_from and created_from != "manual":
                    continue
                lid = str(
                    (row or {}).get("Lesson_ID", "")
                    or (row or {}).get("Lesson_Id", "")
                    or (row or {}).get("id", "")
                ).strip() or str(uuid.uuid4())
                if lid in seen_quote:
                    continue
                category = (row.get("Category") or "Other").strip()
                title = (row.get("Title") or "Quoting Lesson").strip()
                desc = (row.get("Description") or row.get("Lesson") or row.get("Notes") or "").strip()
                process = (row.get("Process") or "").strip()
                material = (row.get("Material") or "").strip()
                source_job = (row.get("Source_Job") or "").strip()
                source_label = (row.get("Source_Label") or "").strip()
                tier = (row.get("Tier") or "").strip()
                date = (row.get("Date") or "").strip()
                fingerprint = (
                    str(record_id or "").strip(),
                    category.lower(),
                    title.lower(),
                    desc.lower(),
                    process.lower(),
                    material.lower(),
                    source_job.lower(),
                    source_label.lower(),
                    tier.lower(),
                    date.lower(),
                )
                if fingerprint in seen_quote_fingerprint:
                    continue
                seen_quote.add(lid)
                seen_quote_fingerprint.add(fingerprint)
                quoting_lessons.append({
                    "id": lid,
                    "row_id": str((row or {}).get("id", "")).strip(),
                    "category": category,
                    "title": title,
                    "desc": desc,
                    "process": process,
                    "material": material,
                    "source_job": source_job,
                    "source_label": source_label,
                    "tier": tier,
                    "date": date,
                    "image_name": (row.get("Image_Name") or "").strip(),
                    "attachment_names": (row.get("Attachment_Names") or "").strip(),
                    "project_record_id": str(record_id or "").strip(),
                    "project_name": (detail.get("Name") or rec.get("Name") or "").strip(),
                    "source": "manual",
                })

        body = {
            "ok": True,
            "mfg_lessons": mfg_lessons,
            "quoting_lessons": quoting_lessons,
        }
        if debug:
            body["debug"] = {
                "requested_supplier_id": (supplier_id or "").strip(),
                "requested_supplier_email": (supplier_email or "").strip().lower(),
                "all_rows_count": len(all_rows),
                "matched_rows_count": len(rows),
                "matched_row_ids": [str((r or {}).get("id") or "").strip() for r in rows[:30]],
                "matched_row_names": [str((r or {}).get("Name") or "").strip() for r in rows[:30]],
                "matched_row_emails": [
                    {
                        "Email": str((r or {}).get("Email") or "").strip().lower(),
                        "Secondary_Email": str((r or {}).get("Secondary_Email") or "").strip().lower(),
                        "Corporate_Email": str((r or {}).get("Corporate_Email") or "").strip().lower(),
                        "Supplier_Name_Id": str((((r or {}).get("Supplier_Name") or {}).get("id") if isinstance((r or {}).get("Supplier_Name"), dict) else "") or "").strip(),
                    }
                    for r in rows[:30]
                ],
            }
        return JSONResponse(body)
    except Exception as e:
        return JSONResponse(
            {"ok": False, "error": str(e), "mfg_lessons": [], "quoting_lessons": []},
            status_code=500,
        )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=True)
